# -*- coding: utf-8 -*-
"""
Genera documentación técnica sobre modelos de IA y algoritmos del SRV-Oratoria-IA
Cubre: modelos abiertos usados, algoritmos implementados y planificados (Sprints 1-3)
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colores ─────────────────────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x1e, 0x29, 0x3b)
AZUL_MEDIO   = RGBColor(0x1e, 0x40, 0xaf)
AZUL_CLARO   = RGBColor(0xdb, 0xe9, 0xfe)
VERDE_BG     = RGBColor(0xdc, 0xfc, 0xe7)
VERDE_TEXT   = RGBColor(0x16, 0x65, 0x34)
AMARILLO_BG  = RGBColor(0xff, 0xf9, 0xc4)
AMARILLO_T   = RGBColor(0x78, 0x35, 0x00)
NARANJA_BG   = RGBColor(0xff, 0xed, 0xcc)
NARANJA_T    = RGBColor(0x92, 0x40, 0x07)
MORADO_BG    = RGBColor(0xed, 0xe9, 0xfe)
MORADO_T     = RGBColor(0x55, 0x21, 0xb0)
GRIS_BG      = RGBColor(0xf1, 0xf5, 0xf9)
GRIS_TEXT    = RGBColor(0x47, 0x55, 0x69)
ROJO_ACENTO  = RGBColor(0xdc, 0x26, 0x26)
BLANCO       = RGBColor(0xff, 0xff, 0xff)

def rgb_hex(rgb) -> str:
    return f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'

# ─── Helpers ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:shd')):
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    tcPr.append(shd)

def set_borders(cell, color='CBD5E1'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def section_title(doc, text, color=AZUL_OSCURO):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BLANCO
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(color))
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.3)
    return p

def subsection(doc, text, color=AZUL_MEDIO):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def subsubsection(doc, text, color=AZUL_OSCURO):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    return p

def body(doc, text, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    return p

def bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p

def make_table(doc, headers, rows, col_widths, estado_col=None, shade_alt=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    # Encabezado
    hrow = t.rows[0]
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hrow.cells[i]
        cell.width = Inches(w)
        set_cell_bg(cell, AZUL_OSCURO)
        set_borders(cell, 'FFFFFF')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = BLANCO
    # Filas
    for ri, row_data in enumerate(rows):
        row = t.add_row()
        for ci, (val, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[ci]
            cell.width = Inches(w)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if shade_alt and ri % 2 == 0:
                set_cell_bg(cell, GRIS_BG)
            set_borders(cell, 'CBD5E1')
            if estado_col is not None and ci == estado_col:
                estado_colors = {
                    'Implementado':  (VERDE_BG,    VERDE_TEXT),
                    'En progreso':   (AMARILLO_BG, AMARILLO_T),
                    'Sprint 2':      (NARANJA_BG,  NARANJA_T),
                    'Sprint 3':      (MORADO_BG,   MORADO_T),
                    'Pendiente':     (GRIS_BG,     GRIS_TEXT),
                }
                bg, fg = estado_colors.get(val, (GRIS_BG, GRIS_TEXT))
                set_cell_bg(cell, bg)
                p2 = cell.paragraphs[0]
                p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p2.add_run(val)
                run.font.size = Pt(8.5)
                run.font.bold = True
                run.font.color.rgb = fg
            else:
                p2 = cell.paragraphs[0]
                run = p2.add_run(val)
                run.font.size = Pt(8.5)
                if ci == 0:
                    run.font.bold = True
                    run.font.color.rgb = AZUL_OSCURO
    return t

def code_block(doc, code, language='python'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1E293B')
    pPr.append(shd)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xa5, 0xf3, 0xfc)
    return p

def info_box(doc, text, bg=AZUL_CLARO, text_color=AZUL_MEDIO):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(bg))
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = text_color
    return p

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ─────────────────────────────────────────────────────────────────────────────
# PORTADA
# ─────────────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\nSISTEMA DE RETROALIMENTACIÓN DE VOCALIZACIÓN\nPARA ESTUDIANTES DE 1° PRIMARIA')
run.font.size  = Pt(20)
run.font.bold  = True
run.font.color.rgb = AZUL_OSCURO

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('SRV-Oratoria-IA\n')
r2.font.size = Pt(15)
r2.font.bold = True
r2.font.color.rgb = AZUL_MEDIO

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run(
    'MODELOS ABIERTOS DE IA Y ALGORITMOS\n'
    'Descripción técnica — Implementados y Planificados\n\n'
)
r3.font.size = Pt(12)
r3.font.color.rgb = GRIS_TEXT

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run(
    'Equipo SRV — UPAO  |  Taller de Tesis S04\n'
    'Ramdhum Arévalo Espinoza (Scrum Master)\n'
    'Arquitectura: FastAPI · Oracle Cloud · Supabase · Vercel\n'
    'Documento generado: 14/05/2026'
)
r4.font.size = Pt(9.5)
r4.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# ÍNDICE MANUAL
# ─────────────────────────────────────────────────────────────────────────────
section_title(doc, 'ÍNDICE DE CONTENIDOS')
indice = [
    ('1.', 'Visión general del pipeline de IA',                 '3'),
    ('2.', 'Tabla resumen de todos los modelos y algoritmos',    '4'),
    ('3.', 'IMPLEMENTADOS — faster-whisper (Whisper medium)',    '5'),
    ('4.', 'IMPLEMENTADOS — parselmouth / Praat (análisis acústico)', '7'),
    ('5.', 'IMPLEMENTADOS — PyAV (conversión de audio)',         '9'),
    ('6.', 'IMPLEMENTADOS — Algoritmo PPM y detección de pausas', '10'),
    ('7.', 'IMPLEMENTADOS — Sistema de retroalimentación por reglas', '12'),
    ('8.', 'SPRINT 2 — spaCy es_core_news_sm (D2 léxico)',       '14'),
    ('9.', 'SPRINT 2 — BETO (bert-base-spanish-wwm-cased) (D2 semántico)', '16'),
    ('10.', 'SPRINT 2 — Proxy acústico de expresividad (D3)',    '18'),
    ('11.', 'SPRINT 2 — Score global integrado D1+D2+D3',        '19'),
    ('12.', 'SPRINT 3 — Claude Haiku (retroalimentación generativa)', '20'),
    ('13.', 'SPRINT 3 — Distancia de Levenshtein (modo lectura)', '22'),
    ('14.', 'Comparativa de modelos: por qué se eligió cada uno', '23'),
    ('15.', 'Infraestructura de despliegue y requisitos de hardware', '24'),
]

t_idx = doc.add_table(rows=0, cols=3)
t_idx.style = 'Table Grid'
for num, titulo, pag in indice:
    row = t_idx.add_row()
    for ci, (val, w) in enumerate(zip([num, titulo, pag], [0.4, 5.5, 0.5])):
        cell = row.cells[ci]
        cell.width = Inches(w)
        p2 = cell.paragraphs[0]
        run = p2.add_run(val)
        run.font.size = Pt(9.5)
        if ci == 0:
            run.font.bold = True
            run.font.color.rgb = AZUL_MEDIO
        if ci == 2:
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run.font.color.rgb = GRIS_TEXT

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. VISIÓN GENERAL
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '1. VISIÓN GENERAL DEL PIPELINE DE IA')
body(doc,
     'El SRV-Oratoria-IA evalúa la oratoria de estudiantes de 1° de primaria mediante '
     'un pipeline de tres dimensiones que combina modelos abiertos de inteligencia artificial '
     'con algoritmos clásicos de procesamiento de señales y lingüística computacional:')

dimensiones = [
    ('D1 — Fluidez oral',
     'Velocidad de habla (PPM), detección de pausas y bloqueos, análisis acústico básico. '
     'Sprint 1: IMPLEMENTADO.',
     VERDE_BG, VERDE_TEXT),
    ('D2 — Coherencia y léxico',
     'Detección de muletillas verbales, riqueza léxica (TTR), coherencia semántica entre ideas. '
     'Sprint 2: EN DESARROLLO.',
     AMARILLO_BG, AMARILLO_T),
    ('D3 — Expresividad vocal',
     'Variación de tono (monotonía), dinámica de intensidad, calidad de voz. '
     'Sprint 2: EN DESARROLLO.',
     NARANJA_BG, NARANJA_T),
]

for dim_name, dim_desc, bg, fg in dimensiones:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(bg))
    pPr.append(shd)
    r1 = p.add_run(f'{dim_name}: ')
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = fg
    r2 = p.add_run(dim_desc)
    r2.font.size = Pt(10)
    r2.font.color.rgb = AZUL_OSCURO

body(doc, '\nFlujo completo de datos en el sistema:')
pipeline_steps = [
    '① Alumno graba voz en el navegador → WaveSurfer.js + RecordPlugin genera blob WebM/Opus',
    '② Frontend POST /audio/analizar con FormData (blob + modo + texto_id opcional)',
    '③ Backend guarda el blob temporal → PyAV convierte WebM → WAV 16kHz mono int16',
    '④ faster-whisper transcribe el WAV → texto + word_timestamps → cálculo PPM + pausas (D1)',
    '⑤ parselmouth analiza el WAV → F0, jitter, shimmer, HNR, intensity (D1 acústico)',
    '⑥ [Sprint 2] spaCy procesa el texto → muletillas, TTR (D2 léxico)',
    '⑦ [Sprint 2] BETO embeddings → cosine similarity entre oraciones → coherencia (D2 semántico)',
    '⑧ [Sprint 2] Praat proxy → variación F0, dinámica → score expresividad (D3)',
    '⑨ Score global = 0.4×D1 + 0.35×D2 + 0.25×D3 → retroalimentación por reglas',
    '⑩ [Sprint 3] Claude Haiku genera consejo personalizado en español simple',
    '⑪ Respuesta JSON → frontend renderiza estrellas, barras, consejos + TTS Web Speech API',
    '⑫ Sesión + métricas guardadas en DB (SQLite/Supabase) → historial y reportes',
]
for step in pipeline_steps:
    bullet(doc, step)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. TABLA RESUMEN
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '2. TABLA RESUMEN DE MODELOS Y ALGORITMOS')

modelos_resumen = [
    ('faster-whisper\n(Whisper medium int8)', 'OpenAI / Systran', 'Reconocimiento automático de voz (ASR)', 'Transcripción + timestamps', 'CPU', 'Implementado'),
    ('parselmouth\n(Praat engine)', 'Paul Boersma / Yannick Jadoul', 'Análisis acústico de voz', 'F0, jitter, shimmer, HNR, intensity', 'CPU', 'Implementado'),
    ('PyAV\n(FFmpeg binding)', 'Comunidad abierta', 'Conversión de formatos de audio/video', 'WAV 16kHz mono int16', 'CPU', 'Implementado'),
    ('Algoritmo PPM\n(propio)', 'Equipo SRV', 'Cálculo de palabras por minuto', 'PPM, duración habla, word_count', 'CPU', 'Implementado'),
    ('Sistema de reglas\n(retroalimentación)', 'Equipo SRV', 'Scoring 1-5 estrellas basado en D1', 'Estrellas, color, mensajes', 'CPU', 'Implementado'),
    ('spaCy\nes_core_news_sm', 'Explosion AI', 'NLP: tokenización, lemas, POS', 'Muletillas, TTR, análisis léxico', 'CPU', 'Sprint 2'),
    ('BETO\n(bert-base-spanish-wwm)', 'DCC UChile', 'LLM BERT fine-tuned en español', 'Embeddings + coherencia semántica', 'CPU/GPU', 'Sprint 2'),
    ('Proxy acústico D3\n(Praat + reglas)', 'Equipo SRV', 'Expresividad vocal por variación F0', 'Score expresividad 0-100', 'CPU', 'Sprint 2'),
    ('Levenshtein\n(difflib Python)', 'Vladimir Levenshtein', 'Distancia de edición entre strings', 'Score fidelidad lectura 0-1', 'CPU', 'Sprint 3'),
    ('Claude Haiku 4.5\n(Anthropic API)', 'Anthropic', 'LLM generativo para consejos adaptativos', 'Texto motivador en español simple', 'Nube', 'Sprint 3'),
]

make_table(doc,
    ['Modelo / Algoritmo', 'Origen', 'Propósito', 'Salida', 'Hardware', 'Estado'],
    modelos_resumen,
    [1.4, 1.2, 1.6, 1.4, 0.65, 0.85],
    estado_col=5
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. FASTER-WHISPER
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '3. IMPLEMENTADO — faster-whisper (Whisper medium int8)')

subsection(doc, '3.1 ¿Qué es?')
body(doc,
     'faster-whisper es una reimplementación del modelo Whisper de OpenAI utilizando el motor '
     'CTranslate2. Permite ejecutar los mismos modelos de reconocimiento de voz de OpenAI con '
     'hasta 4× más velocidad y 2× menos uso de RAM en comparación con la implementación original '
     'de PyTorch, sin necesidad de GPU.')
info_box(doc,
    '⚡ Whisper original (PyTorch): ~8s para audio de 30s | '
    'faster-whisper int8 CPU: ~2s para el mismo audio. '
    'Modelo medium: 769M parámetros, 74 idiomas, WER ~8% en español.',
    VERDE_BG, VERDE_TEXT)

subsection(doc, '3.2 Variante usada: medium int8')
ficha = [
    ('Parámetro',      'Valor'),
    ('Modelo base',    'openai/whisper-medium'),
    ('Versión',        'faster-whisper (Systran/CTranslate2)'),
    ('Cuantización',   'int8 (8-bit enteros)'),
    ('Dispositivo',    'CPU (device="cpu")'),
    ('Idioma forzado', 'es (español)'),
    ('Timestamps',     'word_timestamps=True'),
    ('RAM estimada',   '~1.5 GB en RAM'),
    ('Velocidad',      '< 2s para audios de 30s en Oracle ARM'),
    ('WER español',    '~8% (bajo ruido de fondo)'),
    ('Licencia',       'MIT (Whisper) + Apache 2.0 (CTranslate2)'),
]
make_table(doc, ['Parámetro', 'Valor'], ficha[1:], [2.0, 4.5])

subsection(doc, '3.3 ¿Por qué no PyTorch Whisper?')
razones = [
    'Tamaño: PyTorch Whisper medium ocupa ~3GB RAM vs ~1.5GB con int8.',
    'Velocidad: faster-whisper es 4× más rápido en CPU, crítico para latencia < 2s.',
    'Sin CUDA requerido: Oracle Cloud Free Tier no tiene GPU. faster-whisper int8 corre eficientemente en ARM.',
    'Misma precisión: cuantización int8 pierde < 1% de WER respecto al modelo float32.',
    'Licencia abierta: MIT permite uso comercial y académico sin restricciones.',
]
for r in razones:
    bullet(doc, r)

subsection(doc, '3.4 Cómo se usa en el sistema')
body(doc, 'Archivo: backend/services/dimension1/speech.py')
code_block(doc,
'''from faster_whisper import WhisperModel

model = WhisperModel("medium", device="cpu", compute_type="int8")

def transcribe(wav_path: str) -> dict:
    segments, info = model.transcribe(
        wav_path,
        language="es",
        word_timestamps=True,
        vad_filter=True          # filtra silencios automáticamente
    )
    words = []
    for seg in segments:
        for w in seg.words:
            words.append({
                "word":  w.word,
                "start": w.start,  # segundos
                "end":   w.end
            })
    return {
        "text":   " ".join(w["word"] for w in words),
        "words":  words,
        "duration": info.duration
    }''')

subsection(doc, '3.5 Salidas que produce')
make_table(doc,
    ['Campo', 'Tipo', 'Descripción', 'Ejemplo'],
    [
        ('text',     'str',    'Transcripción completa',                     '"La lluvia en Sevilla"'),
        ('words',    'list',   'Lista de palabras con timestamps start/end', '[{"word":"La","start":0.1,"end":0.4}]'),
        ('duration', 'float',  'Duración total del audio en segundos',       '12.5'),
        ('language', 'str',    'Idioma detectado',                           '"es"'),
    ],
    [1.0, 0.8, 3.0, 2.4]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. PARSELMOUTH / PRAAT
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '4. IMPLEMENTADO — parselmouth / Praat (análisis acústico)')

subsection(doc, '4.1 ¿Qué es?')
body(doc,
     'Praat ("hablar" en neerlandés) es el software de análisis fonético más utilizado en '
     'lingüística clínica y experimental, desarrollado en la Universidad de Ámsterdam. '
     'parselmouth es un binding Python oficial que permite usar todas las funciones de Praat '
     'directamente en código Python sin necesidad de la GUI ni scripts Praat.')
info_box(doc,
    '🔬 Praat es el estándar de oro en análisis de voz académico. '
    'Es usado en patología del lenguaje, lingüística forense y terapia del habla. '
    'Las métricas que extrae son reproducibles y validadas científicamente.',
    AZUL_CLARO, AZUL_MEDIO)

subsection(doc, '4.2 Métricas extraídas')
metricas = [
    ('F0 (frecuencia fundamental)',
     'Tono promedio de la voz. Rango típico: hombres 85-180 Hz, mujeres 165-255 Hz, niños 250-400 Hz.',
     'f0_mean_hz', 'Hz (float)', 'D1 acústico, D3 expresividad'),
    ('F0 std (desviación estándar)',
     'Variación del tono. Valor bajo = monotonía. Valor alto = expresividad vocal.',
     'f0_std_hz', 'Hz (float)', 'D3 expresividad (proxy monotonía)'),
    ('Jitter (variación ciclo a ciclo F0)',
     'Irregularidad en la frecuencia de vibración de cuerdas vocales. Alta = ronquera o disfonía.',
     'jitter_pct', '% (float)', 'D1 calidad de voz'),
    ('Shimmer (variación amplitud ciclo a ciclo)',
     'Irregularidad en la amplitud de vibración. Alta = voz temblorosa o breathy.',
     'shimmer_db', 'dB (float)', 'D1 calidad de voz'),
    ('HNR (Harmonic-to-Noise Ratio)',
     'Relación armónico-ruido. Mide cuánto de la voz es tono puro vs ruido. '
     'HNR > 20 dB = voz limpia. HNR < 10 dB = mucho ruido.',
     'hnr_db', 'dB (float)', 'D1 calidad de voz'),
    ('Intensidad media',
     'Volumen promedio de la voz. Indica si el alumno habla muy suave o muy fuerte.',
     'intensity_mean_db', 'dB (float)', 'D3 expresividad'),
]

make_table(doc,
    ['Métrica', 'Descripción', 'Campo JSON', 'Unidad', 'Dimensión SRV'],
    metricas,
    [1.4, 2.8, 1.1, 0.8, 1.2]
)

subsection(doc, '4.3 Código de análisis acústico')
body(doc, 'Archivo: backend/services/dimension1/prosody.py')
code_block(doc,
'''import parselmouth
from parselmouth.praat import call

def analyze_prosody(wav_path: str) -> dict:
    snd = parselmouth.Sound(wav_path)

    # Pitch (F0)
    pitch  = snd.to_pitch()
    f0_vals = pitch.selected_array["frequency"]
    voiced  = f0_vals[f0_vals > 0]
    f0_mean = round(float(voiced.mean()), 2) if len(voiced) > 0 else None
    f0_std  = round(float(voiced.std()),  2) if len(voiced) > 0 else None

    # Point Process para jitter y shimmer
    pp = call(snd, "To PointProcess (periodic, cc)...", 75, 500)

    # Jitter local (porcentaje)
    jitter = call([snd, pp], "Get jitter (local)...",
                  0, 0, 0.0001, 0.02, 1.3)

    # Shimmer local (dB)
    shimmer = call([snd, pp], "Get shimmer (local_dB)...",
                   0, 0, 0.0001, 0.02, 1.3, 1.6)

    # HNR
    harmonicity = call(snd, "To Harmonicity (cc)...", 0.01, 75, 0.1, 1.0)
    hnr = call(harmonicity, "Get mean...", 0, 0)

    # Intensidad
    intensity = snd.to_intensity()
    intensity_mean = call(intensity, "Get mean...", 0, 0, "energy")

    return {
        "f0_mean_hz":       f0_mean,
        "f0_std_hz":        f0_std,
        "jitter_pct":       round(jitter * 100, 4) if jitter else None,
        "shimmer_db":       round(shimmer, 4) if shimmer else None,
        "hnr_db":           round(hnr, 2) if hnr else None,
        "intensity_mean_db": round(intensity_mean, 2) if intensity_mean else None,
    }''')

subsection(doc, '4.4 Valores de referencia clínicos')
ref_values = [
    ('F0 promedio niño 6-7 años', '250 – 400 Hz', '< 180 Hz o > 500 Hz', 'Problemas de tono'),
    ('Jitter',                    '< 1.04%',       '> 2%',                 'Ronquera / disfonía'),
    ('Shimmer',                   '< 0.46 dB',     '> 1 dB',               'Voz temblorosa'),
    ('HNR',                       '> 20 dB',       '< 10 dB',              'Exceso de ruido/breathy'),
    ('F0 std (expresividad)',      '30 – 80 Hz',    '< 10 Hz',              'Monotonía vocal'),
]
make_table(doc,
    ['Métrica', 'Rango saludable', 'Umbral de alerta', 'Significado clínico'],
    ref_values, [1.6, 1.4, 1.4, 2.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. PyAV
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '5. IMPLEMENTADO — PyAV (conversión de audio WebM → WAV)')

subsection(doc, '5.1 El problema que resuelve')
body(doc,
     'El navegador graba audio con WaveSurfer.js RecordPlugin y genera un blob en formato '
     'WebM/Opus, aunque se nombre con extensión .wav. Praat y faster-whisper requieren WAV '
     'PCM 16kHz mono int16 estándar. Sin esta conversión, Praat lanza:')
info_box(doc,
    '❌  parselmouth.PraatError: Not an audio file.\n'
    'Causa: el archivo tiene cabecera WebM, no cabecera RIFF/WAV.',
    RGBColor(0xfe, 0xe2, 0xe2), RGBColor(0x99, 0x1b, 0x1b))

body(doc, 'PyAV es el binding oficial de FFmpeg para Python. Usa FFmpeg internamente '
          '(una de las bibliotecas multimedia más probadas del mundo) para decodificar '
          'cualquier formato multimedia y recodificar a WAV PCM.')

subsection(doc, '5.2 Código de conversión')
body(doc, 'Archivo: backend/utils/audio.py')
code_block(doc,
'''import av
import numpy as np
import wave

def to_wav(input_path: str, output_path: str) -> None:
    """Convierte cualquier formato de audio a WAV 16kHz mono int16."""
    frames = []
    with av.open(input_path) as container:
        resampler = av.AudioResampler(
            format="s16p",   # PCM 16-bit planar
            layout="mono",   # un canal
            rate=16000       # 16 kHz (requerido por Whisper y Praat)
        )
        for frame in container.decode(audio=0):
            for out_frame in resampler.resample(frame):
                frames.append(out_frame.to_ndarray())
        # Flush final: libera frames retenidos en el buffer del resampler
        for out_frame in resampler.resample(None):
            frames.append(out_frame.to_ndarray())

    audio = np.concatenate(frames, axis=1).flatten().astype("int16")
    with wave.open(output_path, "wb") as wf:
        wf.setnchannels(1)   # mono
        wf.setsampwidth(2)   # 16-bit = 2 bytes
        wf.setframerate(16000)
        wf.writeframes(audio.tobytes())''')

subsection(doc, '5.3 Por qué es crítico el flush final')
body(doc,
     'Los resampleres de FFmpeg tienen un buffer interno. Sin llamar a '
     'resampler.resample(None) al final, los últimos ~100ms del audio '
     'quedan retenidos en el buffer y no se escriben al WAV. Esto causaba '
     'que Whisper "perdiera" el final de la frase del alumno, resultando en '
     'transcripciones incompletas.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. ALGORITMO PPM Y PAUSAS
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '6. IMPLEMENTADO — Algoritmo PPM y detección de pausas')

subsection(doc, '6.1 Cálculo de Palabras Por Minuto (PPM)')
body(doc,
     'El PPM (Palabras Por Minuto) es la métrica central de la Dimensión 1. '
     'Mide la velocidad de habla del alumno excluyendo los silencios para '
     'obtener una medida real del ritmo de elocución.')

body(doc, 'Fórmula:')
code_block(doc,
'''# speech_duration_s = tiempo real hablando (sin pausas largas)
# word_count         = número de palabras reconocidas por Whisper

PPM = (word_count / speech_duration_s) * 60

# Cálculo de speech_duration_s (excluye pausas > 1.5s)
speech_duration_s = 0
for i in range(1, len(words)):
    gap = words[i]["start"] - words[i-1]["end"]
    if gap < 1.5:                          # pausa corta: sí cuenta
        speech_duration_s += gap
    speech_duration_s += (words[i]["end"] - words[i]["start"])  # duración palabra

# Sumar la duración de la primera palabra
if words:
    speech_duration_s += words[0]["end"] - words[0]["start"]''')

info_box(doc,
    '📖 Rango de referencia para niños de 6-7 años (1° primaria):\n'
    '• Muy lento: < 60 PPM  |  Lento: 60-79 PPM  |  Ideal: 80-120 PPM  |  '
    'Rápido: 121-150 PPM  |  Muy rápido: > 150 PPM\n'
    '(Basado en: Rasinski, T.V. "The Fluent Reader", 2010)',
    VERDE_BG, VERDE_TEXT)

subsection(doc, '6.2 Detección de pausas y bloqueos')
body(doc,
     'Una "pausa" es cualquier silencio entre palabras detectado por Whisper. '
     'Un "bloqueo" es una pausa larga (≥ 1.5 segundos) que indica que el alumno '
     'se quedó sin palabras o perdió el hilo del discurso.')

code_block(doc,
'''def analyze_pauses(words: list, long_pause_threshold: float = 1.5) -> dict:
    pauses = []
    for i in range(1, len(words)):
        gap = words[i]["start"] - words[i-1]["end"]
        if gap > 0.3:           # pausa corta (> 300ms)
            pauses.append({
                "start": words[i-1]["end"],
                "end":   words[i]["start"],
                "duration": gap,
                "is_long": gap >= long_pause_threshold
            })

    long_pauses  = sum(1 for p in pauses if p["is_long"])
    total_pauses = len(pauses)
    avg_pause    = sum(p["duration"] for p in pauses) / total_pauses if pauses else 0

    return {
        "total_pauses":  total_pauses,
        "long_pauses":   long_pauses,           # bloqueos
        "avg_pause_s":   round(avg_pause, 2),
        "pauses":        pauses                  # lista completa
    }''')

subsection(doc, '6.3 Tabla de interpretación de resultados D1')
make_table(doc,
    ['Métrica', 'Valor ideal', 'Penalización leve', 'Penalización alta', 'Impacto en estrellas'],
    [
        ('PPM',           '80 – 120',    '60-79 ó 121-150',  '< 60 ó > 150',   '−0 / −1 / −2 pts'),
        ('Bloqueos',      '0',           '1',                '≥ 2',             '−0 / −1 / −2 pts'),
        ('Total pausas',  '< 5',         '5 – 10',           '> 10',            'informativo'),
        ('Jitter',        '< 1.04%',     '1 – 2%',           '> 2%',            'informativo D1'),
        ('HNR',           '> 20 dB',     '10 – 20 dB',       '< 10 dB',         'informativo D1'),
    ],
    [1.2, 1.1, 1.4, 1.4, 1.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. SISTEMA DE RETROALIMENTACIÓN POR REGLAS
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '7. IMPLEMENTADO — Sistema de retroalimentación por reglas (D1)')

subsection(doc, '7.1 Diseño del sistema de scoring')
body(doc,
     'La retroalimentación de D1 usa un sistema de puntos basado en reglas explícitas '
     '(sin aprendizaje automático). Esto permite explicabilidad total: el docente '
     'puede saber exactamente por qué un alumno obtuvo N estrellas.')

body(doc, 'Algoritmo de puntuación:')
code_block(doc,
'''def generate_feedback(ppm_result: dict, pauses_result: dict) -> dict:
    ppm = ppm_result["ppm"]
    bloqueos = pauses_result["long_pauses"]
    score = 0

    # Bloque 1: Puntuación por velocidad (0-3 puntos)
    if   80 <= ppm <= 120:   score += 3   # zona ideal
    elif 70 <= ppm <= 130:   score += 2   # cerca del ideal
    elif 60 <= ppm <= 140:   score += 1   # alejado
    else:                    score += 0   # muy rápido o muy lento

    # Bloque 2: Puntuación por bloqueos (0-2 puntos)
    if   bloqueos == 0:      score += 2   # sin bloqueos
    elif bloqueos == 1:      score += 1   # un bloqueo tolerable
    else:                    score += 0   # dos o más bloqueos

    # Conversión score (0-5) → estrellas (1-5)
    estrellas = max(1, score)

    # Asignación de color y mensaje
    if   score >= 4: color, mensaje = "green",  "¡Excelente! Hablas muy bien"
    elif score >= 3: color, mensaje = "yellow", "¡Bien! Puedes mejorar un poco"
    else:            color, mensaje = "red",    "Sigue practicando, vas a mejorar"

    return {
        "estrellas":         estrellas,
        "color":             color,
        "mensaje_principal": mensaje,
        "detalle_velocidad": _detalle_ppm(ppm),
        "detalle_pausas":    _detalle_pausas(bloqueos),
        "consejos":          _generar_consejos(ppm, bloqueos)
    }''')

subsection(doc, '7.2 Consejos adaptativos por regla')
body(doc,
     'El sistema genera consejos específicos según el problema detectado:')
consejos_tabla = [
    ('PPM > 150', 'Intenta hablar un poco más despacio. Respira entre frases.'),
    ('PPM < 60', 'Habla con más confianza. No tengas miedo de ir un poco más rápido.'),
    ('PPM 60-79', 'Estás casi en el ritmo ideal. Practica con un texto sencillo.'),
    ('Bloqueos ≥ 2', 'Cuando no recuerdes qué decir, haz una pausa corta y sigue.'),
    ('Bloqueos = 1', 'Casi sin bloqueos. Una más de práctica y lo logras.'),
    ('Jitter > 2%', 'Tu voz suena un poco irregular. Relaja la garganta antes de hablar.'),
    ('HNR < 10 dB', 'Hay ruido en el audio. Intenta grabar en un lugar más silencioso.'),
]
make_table(doc, ['Condición detectada', 'Consejo generado'], consejos_tabla, [2.0, 5.0])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. spaCy
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '8. SPRINT 2 — spaCy es_core_news_sm (D2 — análisis léxico)')

subsection(doc, '8.1 ¿Qué es spaCy?')
body(doc,
     'spaCy es una biblioteca de Procesamiento de Lenguaje Natural (NLP) de código abierto '
     'desarrollada por Explosion AI. Es la biblioteca NLP en producción más utilizada '
     'en la industria, con modelos preentrenados para más de 60 idiomas.')
info_box(doc,
    '📦 Modelo a usar: es_core_news_sm (español, small)\n'
    'Tamaño: ~12 MB  |  Velocidad: ~50k palabras/segundo  |  Licencia: MIT\n'
    'Capacidades: tokenización, lematización, POS-tagging, NER, dependencias.',
    NARANJA_BG, NARANJA_T)

subsection(doc, '8.2 Uso 1: Detección de muletillas verbales')
body(doc,
     'Las muletillas (fillers) son palabras o sonidos de relleno que el hablante usa '
     'cuando no sabe qué decir. En español: "este", "emmm", "osea", "bueno", "pues", '
     '"ehhh", "ahhh", "como que", etc. Su presencia excesiva indica dificultad para '
     'organizar el discurso.')

code_block(doc,
'''import spacy
nlp = spacy.load("es_core_news_sm")

MULETILLAS = {
    "este", "esto", "eso", "pues", "bueno", "osea", "o sea",
    "entonces", "digamos", "tipo", "como", "ehh", "emmm",
    "ahhh", "umm", "esteee", "estee", "verdad", "no"
}

def detect_muletillas(texto: str) -> dict:
    doc_nlp = nlp(texto.lower())
    tokens  = [token.text for token in doc_nlp if not token.is_space]
    muletillas_encontradas = []
    conteo_por_tipo = {}

    for token in doc_nlp:
        lemma = token.lemma_.lower()
        texto_token = token.text.lower()
        for m in MULETILLAS:
            if texto_token.startswith(m):   # captura "esteeeee" como "este"
                muletillas_encontradas.append(texto_token)
                conteo_por_tipo[m] = conteo_por_tipo.get(m, 0) + 1

    total = len(tokens)
    count = len(muletillas_encontradas)
    tasa  = round(count / total * 100, 2) if total > 0 else 0

    return {
        "muletillas_count":  count,
        "muletillas_tasa":   tasa,    # porcentaje sobre total palabras
        "por_tipo":          conteo_por_tipo,
        "penaliza":          count > 3  # umbral: más de 3 = penalización en D2
    }''')

subsection(doc, '8.3 Uso 2: Riqueza léxica (Type-Token Ratio)')
body(doc,
     'El TTR (Type-Token Ratio) mide la variedad del vocabulario usado. '
     'Un TTR alto indica que el alumno usa palabras variadas; uno bajo indica '
     'que repite mucho las mismas palabras.')
code_block(doc,
'''def calc_ttr(texto: str) -> dict:
    doc_nlp = nlp(texto.lower())
    # Solo palabras con contenido (excluye stopwords y puntuación)
    tokens = [t.lemma_ for t in doc_nlp
              if not t.is_stop and not t.is_punct and t.is_alpha]
    types  = set(tokens)    # palabras únicas (formas lematizadas)
    ttr    = len(types) / len(tokens) if tokens else 0

    return {
        "word_count":   len(tokens),
        "unique_words": len(types),
        "ttr_score":    round(ttr, 3),   # 0-1 (> 0.4 = bueno para niños)
    }''')

make_table(doc,
    ['TTR', 'Interpretación', 'Para niños 6-7 años'],
    [
        ('> 0.60', 'Vocabulario muy variado',      'Excelente'),
        ('0.40 – 0.60', 'Vocabulario adecuado',   'Bueno'),
        ('0.20 – 0.39', 'Vocabulario limitado',    'En desarrollo'),
        ('< 0.20', 'Vocabulario muy repetitivo',   'Necesita apoyo'),
    ],
    [1.5, 2.5, 2.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. BETO
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '9. SPRINT 2 — BETO (bert-base-spanish-wwm-cased) — coherencia semántica')

subsection(doc, '9.1 ¿Qué es BETO?')
body(doc,
     'BETO es un modelo BERT (Bidirectional Encoder Representations from Transformers) '
     'entrenado exclusivamente en español por el Departamento de Ciencias de la Computación '
     'de la Universidad de Chile (DCC UChile). Utiliza Whole Word Masking (wwm) que mejora '
     'el manejo de morfología española compleja (flexiones verbales, acentos).')
info_box(doc,
    '🤗 Identificador HuggingFace: dccuchile/bert-base-spanish-wwm-cased\n'
    'Tamaño: ~440 MB  |  Parámetros: 110M  |  Vocabulario: 31.002 tokens españoles\n'
    'Entrenado en: Wikipedia ES, SCC, OPUS, multilingual datasets  |  Licencia: MIT',
    MORADO_BG, MORADO_T)

subsection(doc, '9.2 Cómo mide la coherencia semántica')
body(doc,
     'BETO genera embeddings (vectores de alta dimensión) para cada oración. '
     'Oraciones con ideas relacionadas producen embeddings similares (ángulo pequeño). '
     'Oraciones con ideas desconectadas producen embeddings divergentes (ángulo grande). '
     'Usamos cosine similarity entre oraciones consecutivas para medir coherencia local.')

code_block(doc,
'''from transformers import AutoTokenizer, AutoModel
import torch, torch.nn.functional as F

tokenizer = AutoTokenizer.from_pretrained("dccuchile/bert-base-spanish-wwm-cased")
model     = AutoModel.from_pretrained("dccuchile/bert-base-spanish-wwm-cased")

def get_embedding(text: str) -> torch.Tensor:
    inputs  = tokenizer(text, return_tensors="pt",
                        truncation=True, max_length=128)
    with torch.no_grad():
        output = model(**inputs)
    # Mean pooling sobre tokens (excluye [CLS] y [SEP])
    token_embs = output.last_hidden_state[:, 1:-1, :]
    return token_embs.mean(dim=1)

def calc_coherencia(texto: str) -> dict:
    import re
    oraciones = [o.strip() for o in re.split(r"[.!?]+", texto) if len(o.strip()) > 5]
    if len(oraciones) < 2:
        return {"coherencia_score": 1.0, "oraciones": len(oraciones)}

    similarities = []
    for i in range(len(oraciones) - 1):
        emb_a = get_embedding(oraciones[i])
        emb_b = get_embedding(oraciones[i + 1])
        sim   = F.cosine_similarity(emb_a, emb_b).item()
        similarities.append(sim)

    coherencia = sum(similarities) / len(similarities)
    return {
        "coherencia_score":   round(coherencia, 3),   # 0-1 (> 0.6 = bueno)
        "oraciones_analizadas": len(oraciones),
        "similarities":       [round(s, 3) for s in similarities]
    }''')

subsection(doc, '9.3 Interpretación del score de coherencia')
make_table(doc,
    ['Score', 'Significado', 'Mensaje para el alumno'],
    [
        ('0.80 – 1.00', 'Ideas muy conectadas y fluidas',  '¡Tus ideas se entienden muy bien!'),
        ('0.60 – 0.79', 'Coherencia adecuada',             '¡Bien! Tus ideas tienen sentido.'),
        ('0.40 – 0.59', 'Algo desconectado',               'Trata de conectar mejor tus ideas.'),
        ('< 0.40',      'Ideas sin conexión visible',      'Practica organizar lo que vas a decir.'),
    ],
    [1.2, 2.4, 2.8]
)

subsection(doc, '9.4 Consideraciones de rendimiento')
bullet(doc, 'BETO requiere ~440 MB de RAM adicional. Oracle Cloud (24GB) lo soporta holgadamente.')
bullet(doc, 'Para textos cortos (< 5 oraciones) el modelo carga en < 1s tras la primera inferencia.')
bullet(doc, 'Optimización: cargar el modelo al startup (singleton) y no por cada request.')
bullet(doc, 'Alternativa ligera: sentence-transformers con paraphrase-multilingual-MiniLM-L12-v2 (solo ~120MB).')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. PROXY D3
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '10. SPRINT 2 — Proxy acústico de expresividad vocal (D3)')

subsection(doc, '10.1 Concepto')
body(doc,
     'La Dimensión 3 evalúa qué tan "viva" y expresiva suena la voz del alumno. '
     'Una voz monótona (sin variación de tono) y muy suave indica poca confianza o '
     'desinterés. La expresividad no requiere un modelo de IA separado: '
     'se puede calcular como un proxy acústico usando las métricas de Praat que ya extraemos.')

subsection(doc, '10.2 Algoritmo del proxy de expresividad')
code_block(doc,
'''def calc_expresividad(prosodia: dict) -> dict:
    f0_std  = prosodia.get("f0_std_hz",        0) or 0
    f0_mean = prosodia.get("f0_mean_hz",        1) or 1
    hnr_db  = prosodia.get("hnr_db",            0) or 0
    int_db  = prosodia.get("intensity_mean_db", 0) or 0

    # 1. Variación tonal (0-40 pts): coef_variacion = std/mean
    #    Niños: std ~ 40-80 Hz es expresivo. Normalizar sobre 80.
    variacion_tonal = min(40, (f0_std / 80) * 40)

    # 2. Calidad de voz (0-30 pts): HNR > 20 dB es limpio
    calidad_voz = min(30, (hnr_db / 20) * 30) if hnr_db > 0 else 0

    # 3. Volumen adecuado (0-30 pts): 55-75 dB SPL es "voz proyectada"
    if 55 <= int_db <= 75:
        volumen = 30
    elif 45 <= int_db < 55 or 75 < int_db <= 85:
        volumen = 15
    else:
        volumen = 5

    score = round(variacion_tonal + calidad_voz + volumen, 1)  # 0-100
    estrellas = 5 if score >= 80 else 4 if score >= 65 else \\
                3 if score >= 50 else 2 if score >= 35 else 1

    return {
        "expresividad_score": score,
        "estrellas_d3":       estrellas,
        "detalles": {
            "variacion_tonal": round(variacion_tonal, 1),
            "calidad_voz":     round(calidad_voz, 1),
            "volumen":         round(volumen, 1),
        }
    }''')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. SCORE GLOBAL
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '11. SPRINT 2 — Score global integrado D1 + D2 + D3')

body(doc,
     'El sistema combina los tres scores de dimensión en un puntaje global de 0 a 100 '
     'usando pesos ponderados definidos según el énfasis pedagógico de 1° primaria:')

code_block(doc,
'''def calc_score_global(score_d1: float, score_d2: float, score_d3: float) -> dict:
    # Pesos ponderados (suma = 1.0)
    W_D1, W_D2, W_D3 = 0.40, 0.35, 0.25  # D1 prioridad en 1° primaria

    score = round(W_D1 * score_d1 + W_D2 * score_d2 + W_D3 * score_d3, 1)

    if   score >= 85: nivel, color = "Sobresaliente",  "green"
    elif score >= 70: nivel, color = "Bueno",          "green"
    elif score >= 55: nivel, color = "En desarrollo",  "yellow"
    elif score >= 40: nivel, color = "Necesita apoyo", "red"
    else:             nivel, color = "Inicio",         "red"

    return {"score_global": score, "nivel": nivel, "color": color}''')

make_table(doc,
    ['Dimensión', 'Peso', 'Justificación pedagógica'],
    [
        ('D1 — Fluidez oral',      '40%', 'Prioridad máxima en 1° primaria: lectura fluida es el objetivo core del ciclo.'),
        ('D2 — Coherencia y léxico', '35%', 'Organización de ideas y vocabulario adecuados al nivel.'),
        ('D3 — Expresividad vocal',  '25%', 'Importante pero secundaria. Se desarrolla con práctica sostenida.'),
    ],
    [1.5, 0.7, 4.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. CLAUDE HAIKU
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '12. SPRINT 3 — Claude Haiku 4.5 (retroalimentación generativa)')

subsection(doc, '12.1 ¿Qué es Claude Haiku?')
body(doc,
     'Claude Haiku 4.5 es el modelo de lenguaje grande (LLM) más rápido y económico de '
     'Anthropic. Es ideal para tareas de generación de texto corto con alta calidad: '
     'consejos pedagógicos, mensajes motivadores y explicaciones simples para niños. '
     'A diferencia de los modelos locales, Claude Haiku corre en la nube de Anthropic '
     'y se accede via API REST.')
info_box(doc,
    '🤖 Modelo: claude-haiku-4-5-20251001  |  Velocidad: ~200 tokens/s\n'
    'Costo estimado: ~$0.00025 por sesión (25 cents por 1000 sesiones)\n'
    'Contexto: 200k tokens  |  Multilingüe: sí  |  Licencia: API comercial Anthropic',
    MORADO_BG, MORADO_T)

subsection(doc, '12.2 Cómo se integra en el sistema')
body(doc,
     'Claude Haiku recibe las métricas D1+D2+D3 del alumno y genera un consejo '
     'personalizado en español simple, adecuado para niños de 6-7 años. '
     'El prompt está diseñado con técnica few-shot para garantizar respuestas consistentes.')

code_block(doc,
'''import anthropic

client = anthropic.Anthropic()  # usa ANTHROPIC_API_KEY del entorno

SYSTEM_PROMPT = """Eres un asistente educativo amigable que ayuda a niños de 6-7 años
a mejorar su forma de hablar. Tus consejos son:
- Muy cortos (máximo 3 oraciones)
- En español simple (palabras de uso común)
- Motivadores y positivos
- Específicos al problema detectado
Nunca uses tecnicismos como "F0", "jitter" o "PPM"."""

def generar_consejo_ia(metricas: dict) -> str:
    ppm      = metricas["ppm"]["ppm"]
    bloqueos = metricas["pausas"]["long_pauses"]
    estrellas= metricas["retroalimentacion"]["estrellas"]
    muletillas = metricas.get("d2", {}).get("muletillas_count", 0)

    user_msg = f"""El alumno habló a {ppm} palabras por minuto
(ideal: 80-120), tuvo {bloqueos} bloqueos y usó {muletillas} muletillas.
Obtuvo {estrellas} de 5 estrellas.
Escríbele un consejo motivador y específico para mejorar."""

    response = client.messages.create(
        model="claude-haiku-4-5-20251001",
        max_tokens=150,
        system=SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_msg}]
    )
    return response.content[0].text''')

subsection(doc, '12.3 Ejemplos de consejos generados')
make_table(doc,
    ['Situación del alumno', 'Consejo generado por Claude Haiku'],
    [
        ('PPM=155, bloqueos=0, ★★★',
         '¡Muy bien! Hablas con mucha energía. Prueba respirar un poco entre cada idea para que todos puedan seguirte mejor.'),
        ('PPM=52, bloqueos=3, ★★',
         'No te preocupes si a veces te quedas sin palabras, eso le pasa a todos. Practica contando tu día en voz alta cada noche.'),
        ('PPM=95, bloqueos=0, muletillas=8',
         '¡Hablas a la velocidad perfecta! Intenta decir "un momento" en lugar de "este" cuando necesites pensar. ¡Ya casi lo logras!'),
    ],
    [2.5, 5.0]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13. LEVENSHTEIN
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '13. SPRINT 3 — Distancia de Levenshtein (modo lectura)')

subsection(doc, '13.1 ¿Qué es?')
body(doc,
     'La distancia de Levenshtein mide el mínimo número de operaciones de edición '
     '(insertar, eliminar, sustituir) para convertir una cadena en otra. '
     'En el modo lectura, permite comparar lo que el alumno leyó (transcripción Whisper) '
     'con el texto original para detectar omisiones, sustituciones y pronunciaciones incorrectas.')

info_box(doc,
    '📚 Implementación: difflib.SequenceMatcher de Python standard library\n'
    'Sin dependencias externas. Licencia: PSF (Python Software Foundation, libre).\n'
    'Complejidad: O(n²) — eficiente para textos cortos de 1° primaria (< 200 palabras).',
    AZUL_CLARO, AZUL_MEDIO)

code_block(doc,
'''from difflib import SequenceMatcher

def calc_fidelidad_lectura(texto_original: str, transcripcion: str) -> dict:
    # Normalizar: minúsculas, sin signos
    import re
    def normalizar(t):
        return re.sub(r"[^a-záéíóúüñ ]", "", t.lower().strip())

    original   = normalizar(texto_original).split()
    leido      = normalizar(transcripcion).split()

    # Ratio de similitud (0-1): 1 = perfectamente idéntico
    ratio = SequenceMatcher(None, original, leido).ratio()

    # Detalles: palabras correctas, omitidas, sustituidas
    matcher = SequenceMatcher(None, original, leido)
    correctas   = sum(b for _, _, _, _, b in matcher.get_matching_blocks())
    omitidas    = [w for w in original if w not in leido]
    sustituidas = len(original) - correctas - len(omitidas)

    return {
        "fidelidad_score": round(ratio, 3),    # 0-1 (> 0.85 = bueno)
        "palabras_correctas": correctas,
        "palabras_omitidas":  omitidas[:5],     # muestra máx 5
        "palabras_total":     len(original),
    }''')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 14. COMPARATIVA DE MODELOS
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '14. COMPARATIVA: POR QUÉ SE ELIGIÓ CADA MODELO')

body(doc,
     'Para cada tarea de IA se evaluaron alternativas. La siguiente tabla justifica '
     'la elección final considerando: precisión, velocidad, RAM, licencia y costo:')

comparativa = [
    ('ASR (transcripción)',
     'faster-whisper medium int8',
     'Whisper PyTorch, wav2vec2-es, Kaldi, SpeechBrain',
     '4× más rápido que PyTorch Whisper. Sin GPU. ~8% WER. MIT license. 1.5GB RAM.'),
    ('Análisis acústico',
     'parselmouth (Praat)',
     'openSMILE, librosa, STRAIGHT, SpeechBrain',
     'Estándar académico. Métricas clínicas validadas. Sin GPU. 0 MB extra.'),
    ('Conversión audio',
     'PyAV (FFmpeg)',
     'soundfile, pydub, scipy, audioread',
     'Único que maneja WebM/Opus real. FFmpeg soporta 100+ formatos. Sin error Praat.'),
    ('NLP léxico',
     'spaCy es_core_news_sm',
     'NLTK, stanza, polyglot, HuggingFace tokenizers',
     'Más rápido (C++). Modelo pequeño (12MB). Lematización española excelente.'),
    ('Coherencia semántica',
     'BETO (DCC UChile)',
     'multilingual BERT, XLM-R, OpenAI embeddings',
     'Entrenado exclusivamente en español. MIT. Sin costo por token. 440MB.'),
    ('Retroalimentación IA',
     'Claude Haiku 4.5',
     'GPT-4o mini, Gemini Flash, LLaMA 3 local',
     '$0.00025/sesión. 200k contexto. Español excelente. API estable Anthropic.'),
    ('Modo lectura',
     'Levenshtein (difflib stdlib)',
     'BLEU score, ROUGE, sentence-BERT sim',
     'Estándar para edición de texto. Sin deps. Interpretable para docentes.'),
]
make_table(doc,
    ['Tarea', 'Elegido', 'Alternativas evaluadas', 'Justificación de elección'],
    comparativa, [1.0, 1.2, 1.5, 3.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 15. INFRAESTRUCTURA Y HARDWARE
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '15. INFRAESTRUCTURA DE DESPLIEGUE Y REQUISITOS DE HARDWARE')

body(doc,
     'Todos los modelos están seleccionados para funcionar sin GPU (CPU-only), '
     'aprovechando Oracle Cloud Free Tier (4 OCPU ARM Ampere A1 + 24 GB RAM):')

hardware = [
    ('faster-whisper medium int8', 'CPU', '~1.5 GB', '< 2s por audio 30s', 'Oracle Cloud ARM'),
    ('parselmouth / Praat',        'CPU', '~50 MB',  '< 0.5s por audio',   'Oracle Cloud ARM'),
    ('spaCy es_core_news_sm',      'CPU', '~200 MB', '< 0.1s por texto',   'Oracle Cloud ARM'),
    ('BETO',                       'CPU', '~440 MB', '~1-2s por texto',     'Oracle Cloud ARM'),
    ('Proxy D3 (Praat)',           'CPU', '0 extra', '< 0.1s',              'Oracle Cloud ARM'),
    ('Claude Haiku 4.5',           'API', '0 local', '< 0.5s por request',  'Nube Anthropic'),
    ('difflib Levenshtein',        'CPU', '< 1 MB',  '< 0.01s',             'Oracle Cloud ARM'),
    ('TOTAL estimado',             'CPU', '~2.2 GB', '< 5s por sesión',     'Oracle Cloud ARM'),
]
make_table(doc,
    ['Componente', 'Dispositivo', 'RAM estimada', 'Latencia', 'Hosting'],
    hardware, [1.8, 0.9, 1.1, 1.3, 1.6]
)

body(doc, '')
body(doc,
     'Oracle Cloud Free Tier provee 4 OCPU ARM Ampere A1 con 24 GB RAM permanentemente gratis. '
     'El sistema necesita ~2.2 GB de RAM para todos los modelos cargados en memoria, '
     'lo que deja ~21 GB libres para el sistema operativo, PostgreSQL (Supabase) y margen de crecimiento.')

info_box(doc,
    '💡 Estrategia de carga de modelos:\n'
    '• faster-whisper: cargado en startup, singleton global (1 instancia).\n'
    '• BETO: cargado en startup, singleton global (mayor tiempo de carga inicial ~8s).\n'
    '• spaCy: cargado en startup, singleton global.\n'
    '• Praat / parselmouth: no necesita carga; opera por archivo bajo demanda.\n'
    '• Claude Haiku: llamada API solo en Sprint 3, sin carga local.',
    AZUL_CLARO, AZUL_MEDIO)

# ─── Pie ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
p_f = doc.add_paragraph()
p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = p_f.add_run(
    '────────────────────────────────────────────────────────────────────\n'
    'SRV-Oratoria-IA · Taller de Tesis S04 · UPAO 2026\n'
    'Ramdhum Arévalo Espinoza — Scrum Master\n'
    'Modelos: faster-whisper · parselmouth · spaCy · BETO · Claude Haiku · Levenshtein'
)
rf.font.size = Pt(8)
rf.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

# ─── Guardar ─────────────────────────────────────────────────────────────────
output = (r'c:\Users\simpl\Documents\Proyecto - Informe - Taller y Tesis'
          r'\SRV-Oratoria-IA_Modelos_Algoritmos.docx')
doc.save(output)
print(f'[OK] Documento generado: {output}')
print('     Secciones: 15')
print('     Modelos documentados: 7 (implementados + planificados)')
print('     Algoritmos: PPM, pausas, scoring por reglas, proxy D3, score global, Levenshtein')

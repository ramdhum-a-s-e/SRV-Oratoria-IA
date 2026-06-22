# -*- coding: utf-8 -*-
"""
Genera el documento Word de documentacion tecnica del proyecto SRV-Oratoria-IA.
Ejecutar desde la raiz del repositorio:
    python generar_documentacion.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos globales ──────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def set_heading(paragraph, level, text, color=None):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
    run.font.name = 'Calibri'
    run.font.bold = True
    sizes = {1: 18, 2: 15, 3: 13}
    run.font.size = Pt(sizes.get(level, 12))
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(text):
    p = doc.add_heading(text, level=1)
    set_heading(p, 1, text, (31, 73, 125))
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    set_heading(p, 2, text, (21, 96, 130))
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    set_heading(p, 3, text, (43, 145, 175))
    return p

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F2F2F2')
    shading.set(qn('w:val'), 'clear')
    p._p.get_or_add_pPr().append(shading)
    return p

def tabla(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2156AE')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        for run in hdr[i].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(10)
    doc.add_paragraph()
    return t


# ═══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('UNIVERSIDAD PRIVADA ANTENOR ORREGO')
run.font.bold = True
run.font.size = Pt(14)
run.font.name = 'Calibri'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SRV - SISTEMA DE RETROALIMENTACION POR VOZ')
run.font.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(31, 73, 125)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documentacion Tecnica del Sistema')
run.font.size = Pt(13)
run.font.name = 'Calibri'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Taller Integrador 1 - Ingenieria de Sistemas')
run.font.size = Pt(11)
run.font.name = 'Calibri'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Fecha: {datetime.date.today().strftime("%d de %B de %Y")}')
run.font.size = Pt(11)
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Equipo: Anthony Lezcano Saavedra | Ramdhum Arevalo Espinoza')
run.font.size = Pt(11)
run.font.name = 'Calibri'

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. DESCRIPCION GENERAL
# ═══════════════════════════════════════════════════════════════════════════════
h1('1. Descripcion General del Proyecto')

para(
    'El SRV (Sistema de Retroalimentacion por Voz) es una aplicacion web que evalua '
    'la fluidez oral de alumnos de 1er grado de primaria de la I.E. Juan Jose Farfan, '
    'Lancones, Piura. El sistema graba la voz del alumno, la analiza con modelos de '
    'Inteligencia Artificial y entrega retroalimentacion inmediata sobre su desempeno.'
)

h2('1.1 Objetivos del Sistema')
bullet('Medir la velocidad de habla en Palabras por Minuto (PPM) con un objetivo de 80-120 PPM para 1er grado.')
bullet('Detectar pausas y bloqueos durante la lectura o expresion oral.')
bullet('Analizar la calidad de la voz (tono, jitter, shimmer, HNR) con tecnologia Praat.')
bullet('Generar retroalimentacion automatica comprensible para ninos de 6-7 anos.')
bullet('Registrar el historial de practicas para seguimiento de la evolucion del alumno.')

h2('1.2 Modos de Practica')
bullet('Lectura guiada: el sistema propone un texto y el alumno lo lee en voz alta.')
bullet('Expresion oral libre: el alumno habla sobre cualquier tema de su eleccion.')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 2. ARQUITECTURA
# ═══════════════════════════════════════════════════════════════════════════════
h1('2. Arquitectura del Sistema')

para(
    'El sistema sigue una arquitectura cliente-servidor de tres capas: frontend (React), '
    'backend (FastAPI + Python) y base de datos (PostgreSQL). Para produccion se despliega '
    'en Oracle Cloud Free Tier (backend), Vercel (frontend) y Supabase (base de datos), '
    'con costo operativo de S/ 0.'
)

h2('2.1 Diagrama de Arquitectura')
code_block(
    'ALUMNO (navegador web)\n'
    '        |\n'
    '        | HTTPS\n'
    '        |\n'
    '   +-----------+         +------------------------+\n'
    '   |  VERCEL   |  API    |  ORACLE CLOUD ARM      |\n'
    '   | (Frontend)|<------->|  (Backend FastAPI)     |\n'
    '   |  React    |         |  Whisper + Praat       |\n'
    '   +-----------+         +----------+-------------+\n'
    '                                    |\n'
    '                                    | PostgreSQL\n'
    '                                    |\n'
    '                            +-------+--------+\n'
    '                            |   SUPABASE     |\n'
    '                            |  (Base datos)  |\n'
    '                            +----------------+\n'
)

h2('2.2 Stack Tecnologico')
tabla(
    ['Componente', 'Tecnologia', 'Version', 'Proposito'],
    [
        ('Backend', 'FastAPI + Python', '3.13.7', 'API REST, logica de negocio'),
        ('ASR', 'faster-whisper', 'ctranslate2', 'Transcripcion de voz a texto'),
        ('Prosodia', 'parselmouth / Praat', '0.4.5', 'Analisis acustico de la voz'),
        ('Conversion audio', 'PyAV', 'bundled', 'WebM/Opus a WAV 16kHz mono'),
        ('Frontend', 'React + Vite', '18 / 6', 'Interfaz de usuario'),
        ('Grabacion', 'WaveSurfer.js', '7.x', 'Waveform + grabacion en navegador'),
        ('Base de datos', 'PostgreSQL / SQLite', '15 / 3', 'Almacenamiento de datos'),
        ('ORM', 'SQLAlchemy', '2.x', 'Mapeo objeto-relacional'),
        ('Auth', 'JWT + Argon2', 'python-jose', 'Autenticacion de usuarios'),
        ('Despliegue BD', 'Supabase', 'Free tier', 'PostgreSQL gestionado gratis'),
        ('Despliegue Backend', 'Oracle Cloud ARM', 'Free tier', '4 OCPU + 24GB RAM gratis'),
        ('Despliegue Frontend', 'Vercel', 'Free tier', 'Hosting estatico gratis'),
    ]
)

h2('2.3 Flujo de una Sesion de Practica')
bullet('1. El alumno inicia sesion con usuario y contrasena (JWT).')
bullet('2. Elige el modo: Lectura o Expresion Libre.')
bullet('3. Si elige Lectura, el sistema muestra un texto propuesto.')
bullet('4. El alumno presiona "Empezar a hablar" y graba con el microfono.')
bullet('5. Al presionar "Listo, analizar", el audio (WebM) se envia al backend.')
bullet('6. El backend convierte WebM a WAV 16kHz mono con PyAV.')
bullet('7. faster-whisper transcribe el audio y extrae timestamps por palabra.')
bullet('8. Se calculan: PPM, pausas, bloqueos.')
bullet('9. Praat analiza la calidad acustica de la voz (F0, jitter, shimmer, HNR).')
bullet('10. Se genera retroalimentacion automatica con reglas (estrellas 1-5, mensajes, consejos).')
bullet('11. Los resultados se guardan en la base de datos (sesion + metricas).')
bullet('12. El frontend muestra la retroalimentacion en pantalla y la lee en voz alta (TTS).')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 3. ESTRUCTURA DEL PROYECTO
# ═══════════════════════════════════════════════════════════════════════════════
h1('3. Estructura de Archivos del Proyecto')

code_block(
    'SRV-Oratoria-IA/\n'
    '├── backend/\n'
    '│   ├── main.py                  # Punto de entrada FastAPI\n'
    '│   ├── config.py                # Variables de entorno\n'
    '│   ├── database.py              # Conexion SQLAlchemy\n'
    '│   ├── validate_d1.py           # Script de pruebas D1 por terminal\n'
    '│   ├── requirements.txt         # Dependencias Python\n'
    '│   ├── .env.example             # Variables de entorno de ejemplo\n'
    '│   ├── models/\n'
    '│   │   ├── user.py              # Modelo Usuario\n'
    '│   │   ├── session.py           # Modelos Sesion y TextoLectura\n'
    '│   │   └── metrics.py           # Modelo ResultadoD1\n'
    '│   ├── routers/\n'
    '│   │   ├── auth.py              # Endpoints de autenticacion\n'
    '│   │   ├── audio.py             # Endpoints de analisis de voz\n'
    '│   │   └── metrics.py           # Endpoints de historial\n'
    '│   ├── services/\n'
    '│   │   ├── audio_processor.py   # Singleton de modelos Whisper\n'
    '│   │   ├── dimension1/\n'
    '│   │   │   ├── fluency.py       # Transcripcion y PPM\n'
    '│   │   │   ├── pauses.py        # Deteccion de pausas\n'
    '│   │   │   ├── prosody.py       # Analisis Praat\n'
    '│   │   │   └── feedback.py      # Generacion de retroalimentacion\n'
    '│   │   ├── dimension2/          # (pendiente) Claridad y coherencia\n'
    '│   │   └── dimension3/          # (pendiente) Seguridad emocional\n'
    '│   ├── schemas/\n'
    '│   │   ├── auth.py              # Esquemas Pydantic de auth\n'
    '│   │   ├── audio.py             # Esquemas de textos de lectura\n'
    '│   │   └── metrics.py           # Esquemas de historial\n'
    '│   └── utils/\n'
    '│       ├── auth.py              # JWT, hash, dependencia usuario\n'
    '│       └── audio.py             # Conversion de formato de audio\n'
    '└── frontend/\n'
    '    └── src/\n'
    '        ├── App.jsx              # Router principal (react-router-dom)\n'
    '        ├── api.js               # Cliente axios con JWT automatico\n'
    '        ├── context/\n'
    '        │   └── AuthContext.jsx  # Estado global de autenticacion\n'
    '        └── pages/\n'
    '            ├── Login.jsx        # Pagina de login y registro\n'
    '            ├── ModoSeleccion.jsx # Eleccion de modo de practica\n'
    '            ├── Practica.jsx     # Grabacion + retroalimentacion\n'
    '            └── Historial.jsx    # Historial y graficos\n'
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 4. BASE DE DATOS
# ═══════════════════════════════════════════════════════════════════════════════
h1('4. Base de Datos')

para(
    'El sistema utiliza SQLAlchemy como ORM. En desarrollo local se usa SQLite '
    '(archivo srv_dev.db generado automaticamente). En produccion se usa PostgreSQL '
    'en Supabase. La base de datos se crea automaticamente al iniciar el servidor.'
)

h2('4.1 Modelo Entidad-Relacion')
code_block(
    'usuarios\n'
    '  id, nombre, apellido, username, password_hash, rol, grado, seccion, created_at\n'
    '       |\n'
    '       | 1:N\n'
    '       |\n'
    '   sesiones\n'
    '     id, usuario_id, modo (lectura|libre), texto_id, created_at\n'
    '          |                                    |\n'
    '          | 1:1                            N:1 |\n'
    '          |                                    |\n'
    '   resultados_d1                       textos_lectura\n'
    '     id, sesion_id, transcripcion,       id, titulo, contenido, nivel\n'
    '     ppm, word_count, speech_duration_s,\n'
    '     total_pauses, long_pauses, avg_pause_s,\n'
    '     f0_mean_hz, f0_std_hz, jitter_pct,\n'
    '     shimmer_db, hnr_db, intensity_mean_db,\n'
    '     estrellas, feedback_json\n'
)

h2('4.2 Descripcion de Tablas')
tabla(
    ['Tabla', 'Descripcion', 'Registros esperados'],
    [
        ('usuarios', 'Alumnos y docentes registrados en el sistema', '25-30 alumnos por aula'),
        ('sesiones', 'Cada practica de voz realizada por un alumno', 'Varias por alumno por semana'),
        ('resultados_d1', 'Metricas de fluidez de cada sesion (1:1 con sesion)', 'Igual que sesiones'),
        ('textos_lectura', 'Textos propuestos para el modo lectura', '3 textos iniciales (seed)'),
    ]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 5. BACKEND - MODULOS PRINCIPALES
# ═══════════════════════════════════════════════════════════════════════════════
h1('5. Backend - Modulos Principales')

h2('5.1 main.py — Punto de Entrada')
para(
    'Inicializa la aplicacion FastAPI, registra los routers, configura CORS para '
    'permitir peticiones desde el frontend, crea las tablas de la base de datos al '
    'iniciar y carga los textos de lectura iniciales (seed).'
)
bullet('GET /  →  health check, retorna estado del servidor.')
bullet('Routers registrados: /auth, /audio, /metrics.')

h2('5.2 config.py — Configuracion')
para(
    'Lee las variables de entorno del archivo .env. Si no existen, usa valores por '
    'defecto seguros para desarrollo local.'
)
tabla(
    ['Variable', 'Valor por defecto', 'Descripcion'],
    [
        ('WHISPER_DEVICE', 'cpu', 'Dispositivo de inferencia (cpu o cuda)'),
        ('WHISPER_MODEL_FINAL', 'medium', 'Modelo Whisper para analisis final'),
        ('WHISPER_COMPUTE_TYPE', 'int8', 'Precision numerica (int8 = mas rapido)'),
        ('DATABASE_URL', 'sqlite:///./srv_dev.db', 'Cadena de conexion a la base de datos'),
        ('JWT_SECRET', 'dev_secret_...', 'Clave secreta para firmar tokens JWT'),
        ('JWT_EXPIRE_MINUTES', '1440', 'Duracion del token (24 horas)'),
    ]
)

h2('5.3 services/dimension1/ — Motor de Evaluacion D1')
para(
    'Contiene toda la logica de evaluacion de Fluidez Mecanica y Prosodia. '
    'Es el nucleo del sistema.'
)

h3('fluency.py')
para('Transcripcion de audio y calculo de PPM.')
bullet('transcribe(audio_path, model): usa faster-whisper para transcribir el audio con timestamps por palabra. Retorna lista de WordToken (palabra, inicio, fin) y el texto completo.')
bullet('calculate_ppm(words): calcula Palabras por Minuto dividiendo el numero de palabras entre la duracion de habla. Objetivo para 1er grado: 80-120 PPM.')

h3('pauses.py')
para('Deteccion de pausas a partir de los timestamps de las palabras.')
bullet('detect_pauses(words): analiza los intervalos entre palabras consecutivas. Una pausa >= 0.5s se registra. Una pausa >= 2.0s se considera bloqueo. Retorna conteo total, bloqueos y promedio.')

h3('prosody.py')
para('Analisis acustico con la libreria parselmouth (interfaz Python de Praat).')
bullet('analyze_prosody(audio_path): analiza el archivo WAV y extrae F0 (tono fundamental), jitter (irregularidad de periodo), shimmer (irregularidad de amplitud), HNR (relacion armonicos/ruido) e intensidad promedio.')
bullet('IMPORTANTE: jitter se calcula con call(pp, ...) - solo PointProcess. Shimmer necesita call([snd, pp], ...).')

h3('feedback.py')
para('Generacion de retroalimentacion basada en reglas.')
bullet('generate_feedback(ppm_result, pauses_result): evalua las metricas y asigna 1-5 estrellas segun el rango de PPM (80-120 = ideal) y numero de bloqueos (0 = perfecto). Genera mensajes en lenguaje simple para ninos y consejos de mejora.')

h2('5.4 utils/audio.py — Conversion de Audio')
para(
    'El navegador graba audio en formato WebM/Opus. Praat solo acepta WAV. '
    'Esta utilidad convierte cualquier formato a WAV 16kHz mono usando PyAV '
    '(que incluye ffmpeg internamente, sin dependencias externas).'
)
bullet('to_wav(input_path, output_path): abre el archivo con PyAV, resamplea a 16kHz mono y guarda como WAV PCM 16-bit.')

h2('5.5 utils/auth.py — Autenticacion')
para('Manejo de contrasenas y tokens JWT.')
bullet('hash_password(password): genera hash Argon2 de la contrasena.')
bullet('verify_password(plain, hashed): verifica contrasena contra el hash.')
bullet('create_token(user_id, rol): genera token JWT con expiracion configurable.')
bullet('get_current_user(token, db): dependencia FastAPI que valida el token y retorna el usuario. Se usa en endpoints protegidos.')

h2('5.6 routers/auth.py — Endpoints de Autenticacion')
tabla(
    ['Metodo', 'Ruta', 'Descripcion', 'Auth requerida'],
    [
        ('POST', '/auth/register', 'Registra nuevo usuario, retorna token JWT', 'No'),
        ('POST', '/auth/login', 'Inicia sesion, retorna token JWT', 'No'),
        ('GET', '/auth/me', 'Retorna datos del usuario autenticado', 'Si'),
    ]
)

h2('5.7 routers/audio.py — Endpoints de Analisis de Voz')
tabla(
    ['Metodo', 'Ruta', 'Descripcion', 'Auth requerida'],
    [
        ('GET', '/audio/textos', 'Lista textos de lectura disponibles', 'No'),
        ('POST', '/audio/analizar', 'Recibe audio, analiza D1, guarda en DB, retorna resultados', 'Si'),
    ]
)
para('El endpoint POST /audio/analizar recibe: file (audio), modo (lectura|libre), texto_id (opcional).')

h2('5.8 routers/metrics.py — Endpoints de Historial')
tabla(
    ['Metodo', 'Ruta', 'Descripcion', 'Auth requerida'],
    [
        ('GET', '/metrics/historial', 'Retorna ultimas 20 sesiones del usuario con sus estrellas y PPM', 'Si'),
        ('GET', '/metrics/sesion/{id}', 'Retorna el resultado completo de una sesion especifica', 'Si'),
    ]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 6. FRONTEND - PAGINAS
# ═══════════════════════════════════════════════════════════════════════════════
h1('6. Frontend - Paginas y Navegacion')

h2('6.1 Estructura de Navegacion')
code_block(
    '/login         -> Login.jsx      (publica, redirige a /modos si ya hay sesion)\n'
    '/modos         -> ModoSeleccion  (privada, elige tipo de practica)\n'
    '/practica      -> Practica.jsx   (privada, grabacion + retroalimentacion)\n'
    '/historial     -> Historial.jsx  (privada, historial y graficos)\n'
    '\n'
    'Rutas no existentes -> redirigen a /login\n'
)

h2('6.2 AuthContext.jsx — Estado Global de Autenticacion')
para(
    'Guarda el token JWT y los datos del usuario en localStorage para que persistan '
    'al recargar la pagina. Provee las funciones login() y logout() a todos los componentes.'
)

h2('6.3 api.js — Cliente HTTP')
para(
    'Instancia de axios configurada con la URL base del backend. Interceptor automatico '
    'que agrega el header Authorization: Bearer <token> a todas las peticiones si hay sesion activa.'
)

h2('6.4 Login.jsx')
para(
    'Formulario con dos modos: "Iniciar sesion" y "Registrarse". '
    'Llama a POST /auth/login o POST /auth/register y guarda el token en AuthContext.'
)

h2('6.5 ModoSeleccion.jsx')
para(
    'Muestra dos opciones: Lectura y Expresion Libre. Si elige Lectura, '
    'muestra la lista de textos disponibles (GET /audio/textos) y el alumno selecciona uno. '
    'Navega a /practica con parametros modo y texto_id en la URL.'
)

h2('6.6 Practica.jsx')
para(
    'Pagina principal de practica. Muestra el texto a leer (si es modo lectura), '
    'el visualizador de forma de onda (WaveSurfer.js), el boton de grabacion y los resultados.'
)
bullet('Al iniciar: crea instancia WaveSurfer + RecordPlugin y activa el microfono.')
bullet('Al detener: envia el audio a POST /audio/analizar con FormData (archivo + modo + texto_id).')
bullet('Al recibir respuesta: muestra panel de retroalimentacion y lee el resultado en voz alta con Web Speech API.')
bullet('Panel de retroalimentacion: estrellas, barra de velocidad PPM, contadores de bloqueos, consejos, seccion colapsable con datos tecnicos para el docente.')

h2('6.7 Historial.jsx')
para(
    'Muestra el resumen de todas las practicas del alumno autenticado.'
)
bullet('Tarjetas de resumen: total de practicas, promedio de estrellas, promedio de bloqueos.')
bullet('Grafico de barras de PPM por sesion (barras de colores segun si esta en rango ideal).')
bullet('Lista de sesiones con fecha, modo, estrellas y PPM.')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 7. METRICAS DE EVALUACION D1
# ═══════════════════════════════════════════════════════════════════════════════
h1('7. Metricas de Evaluacion - Dimension 1')

h2('7.1 Palabras por Minuto (PPM)')
para('Mide la velocidad de habla del alumno.')
tabla(
    ['Rango PPM', 'Evaluacion', 'Mensaje al alumno'],
    [
        ('< 60', 'Muy lento', 'Hablas muy despacio. Intenta leer con mas energia.'),
        ('60 - 79', 'Lento', 'Hablas un poco despacio. Casi llegas al ritmo ideal.'),
        ('80 - 120', 'Ideal', 'Hablas a una velocidad perfecta.'),
        ('121 - 140', 'Rapido', 'Hablas un poco rapido. Respira y ve mas despacio.'),
        ('> 140', 'Muy rapido', 'Hablas muy rapido. Es dificil entenderte.'),
    ]
)

h2('7.2 Pausas y Bloqueos')
tabla(
    ['Metrica', 'Umbral', 'Descripcion'],
    [
        ('Pausa detectable', '>= 0.5 segundos', 'Silencio entre palabras considerado pausa'),
        ('Bloqueo', '>= 2.0 segundos', 'Pausa larga que indica nerviosismo o olvido'),
    ]
)

h2('7.3 Metricas Prosodicas (Praat)')
tabla(
    ['Metrica', 'Descripcion', 'Umbral de normalidad'],
    [
        ('F0 promedio (Hz)', 'Tono fundamental de la voz', 'Variable por persona'),
        ('F0 std (Hz)', 'Variacion del tono - indica expresividad', 'Mayor valor = mas expresivo'),
        ('Jitter (%)', 'Irregularidad en la frecuencia de vibracion', '< 1.04% = voz sana'),
        ('Shimmer (dB)', 'Irregularidad en la amplitud de vibracion', '< 0.35 dB = estable'),
        ('HNR (dB)', 'Relacion armonicos/ruido - claridad de la voz', '> 20 dB = voz clara'),
        ('Intensidad (dB)', 'Volumen promedio de la voz', 'Referencial'),
    ]
)

h2('7.4 Sistema de Estrellas')
para('La retroalimentacion principal se expresa en 1 a 5 estrellas calculadas asi:')
bullet('PPM en rango 80-120: 3 puntos. En rango cercano: 2 puntos. Fuera de rango: 1 punto.')
bullet('Sin bloqueos: 2 puntos. 1 bloqueo: 1 punto. 2+ bloqueos: 0 puntos.')
bullet('Total = suma de puntos, ajustado a rango 1-5 estrellas.')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 8. INSTALACION Y EJECUCION
# ═══════════════════════════════════════════════════════════════════════════════
h1('8. Instalacion y Ejecucion Local')

h2('8.1 Requisitos Previos')
bullet('Python 3.11 o superior')
bullet('Node.js 18 o superior (npm incluido)')
bullet('Git')

h2('8.2 Clonar el Repositorio')
code_block('git clone https://github.com/ALS-12321/SRV-Oratoria-IA.git\ncd SRV-Oratoria-IA')

h2('8.3 Configurar el Backend')
code_block(
    'cd backend\n'
    'python -m venv venv\n'
    '.\\venv\\Scripts\\activate          # Windows\n'
    'pip install -r requirements.txt\n'
)

h2('8.4 Levantar el Backend')
code_block(
    'cd backend\n'
    '$env:PYTHONIOENCODING="utf-8"\n'
    '.\\venv\\Scripts\\uvicorn.exe main:app --reload\n'
    '# Servidor disponible en: http://127.0.0.1:8000\n'
    '# Documentacion API en:   http://127.0.0.1:8000/docs\n'
)

h2('8.5 Levantar el Frontend')
code_block(
    'cd frontend\n'
    'npm install\n'
    'npm run dev\n'
    '# Aplicacion disponible en: http://localhost:5173\n'
)

h2('8.6 Primer Uso')
bullet('1. Abrir http://localhost:5173 en el navegador.')
bullet('2. Registrar una cuenta nueva con nombre, apellido, usuario y contrasena.')
bullet('3. Elegir modo de practica: Lectura o Expresion Libre.')
bullet('4. Grabar y recibir retroalimentacion.')
bullet('5. Ver el historial de practicas en la seccion "Mi historial".')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 9. API REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
h1('9. Referencia de API')

para('Todos los endpoints estan documentados automaticamente en http://127.0.0.1:8000/docs (Swagger UI).')

h2('9.1 Ejemplo: Registro de Usuario')
code_block(
    'POST /auth/register\n'
    'Content-Type: application/json\n\n'
    '{\n'
    '  "nombre": "Juan",\n'
    '  "apellido": "Perez",\n'
    '  "username": "juan.perez",\n'
    '  "password": "mi_contrasena",\n'
    '  "grado": "1ro",\n'
    '  "seccion": "A"\n'
    '}\n\n'
    'Respuesta:\n'
    '{\n'
    '  "access_token": "eyJ...",\n'
    '  "token_type": "bearer",\n'
    '  "user": { "id": 1, "nombre": "Juan", ... }\n'
    '}\n'
)

h2('9.2 Ejemplo: Analizar Audio')
code_block(
    'POST /audio/analizar\n'
    'Authorization: Bearer <token>\n'
    'Content-Type: multipart/form-data\n\n'
    'file: <archivo_audio.wav>\n'
    'modo: "lectura"\n'
    'texto_id: 1\n\n'
    'Respuesta:\n'
    '{\n'
    '  "sesion_id": 5,\n'
    '  "transcripcion": "Habia una vez un perrito...",\n'
    '  "ppm": { "ppm": 95.3, "word_count": 42, "speech_duration_s": 26.4 },\n'
    '  "pausas": { "total_pauses": 2, "long_pauses": 0, "avg_pause_s": 0.6 },\n'
    '  "prosodia": { "f0_mean_hz": 185.2, "jitter_pct": 0.45, ... },\n'
    '  "retroalimentacion": {\n'
    '    "estrellas": 4,\n'
    '    "color": "green",\n'
    '    "mensaje_principal": "Muy bien hecho!",\n'
    '    "detalle_velocidad": "Hablas a una velocidad perfecta.",\n'
    '    "consejos": []\n'
    '  }\n'
    '}\n'
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 10. PENDIENTES Y PROXIMAS VERSIONES
# ═══════════════════════════════════════════════════════════════════════════════
h1('10. Trabajo Pendiente y Proximas Versiones')

h2('10.1 Mejoras en D1 (Sprint actual)')
bullet('Corregir calculo de PPM para excluir tiempo de pausas largas.')
bullet('Detectar muletillas prolongadas ("esteee", "ahhhh") como bloqueos aunque Whisper las transcriba como palabras.')

h2('10.2 Dimension 2 - Claridad y Coherencia (proximo sprint)')
bullet('Deteccion de muletillas lexicas ("este", "o sea", "em") con spaCy.')
bullet('Type-Token Ratio (TTR) para medir riqueza lexica.')
bullet('Analisis de coherencia semantica entre oraciones con modelo BETO (BERT en espanol).')

h2('10.3 Dimension 3 - Seguridad Emocional (proximo sprint)')
bullet('Proxy acustico de nerviosismo usando parselmouth: variacion de F0 en el tiempo, variacion de energia y velocidad por segmento.')
bullet('Sin modelos pesados de reconocimiento de emociones (limitacion declarada en la tesis).')

h2('10.4 Funcionalidades Pendientes')
bullet('Reporte PDF al final de cada sesion.')
bullet('Panel de docente para ver el progreso de todos los alumnos del aula.')
bullet('Despliegue en Oracle Cloud + Vercel + Supabase.')
bullet('Reduccion de ruido como preprocesamiento antes de Whisper.')

# ═══════════════════════════════════════════════════════════════════════════════
# Guardar
# ═══════════════════════════════════════════════════════════════════════════════
output = 'SRV-Oratoria-IA_Documentacion_Tecnica.docx'
doc.save(output)
print(f'Documento generado: {output}')

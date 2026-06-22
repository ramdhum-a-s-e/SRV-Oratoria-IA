# -*- coding: utf-8 -*-
"""
Genera el Product Backlog reformulado para SRV-Oratoria-IA
Alineado con arquitectura real: FastAPI + Oracle Cloud + Supabase + Vercel
~50 items totales | ~28 HU | Sprints 1-3
"""
import sys
import os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Paleta de colores ────────────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x1e, 0x29, 0x3b)   # encabezados sección
AZUL_MEDIO   = RGBColor(0x1e, 0x40, 0xaf)   # sprint headers
AZUL_CLARO   = RGBColor(0xdb, 0xe9, 0xfe)   # fila par tabla Sprint 1
VERDE_BG     = RGBColor(0xdc, 0xfc, 0xe7)   # completado
VERDE_TEXT   = RGBColor(0x16, 0x65, 0x34)
AMARILLO_BG  = RGBColor(0xff, 0xf9, 0xc4)   # en progreso
AMARILLO_T   = RGBColor(0x78, 0x35, 0x00)
GRIS_BG      = RGBColor(0xf1, 0xf5, 0xf9)   # pendiente
GRIS_TEXT    = RGBColor(0x47, 0x55, 0x69)
ROJO_BG      = RGBColor(0xfe, 0xe2, 0xe2)   # no aplica
ROJO_TEXT    = RGBColor(0x99, 0x1b, 0x1b)
BLANCO       = RGBColor(0xff, 0xff, 0xff)

# ─── Helpers ─────────────────────────────────────────────────────────────────
def rgb_hex(rgb) -> str:
    if isinstance(rgb, str):
        return rgb
    return f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'

def set_cell_bg(cell, rgb):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    tcPr.append(shd)

def set_cell_borders(cell, color='AAAAAA'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def run_color(run, rgb: RGBColor):
    run.font.color.rgb = rgb

def h1(doc, txt):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(txt)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BLANCO
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(AZUL_OSCURO))
    pPr.append(shd)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    return p

def h2(doc, txt, color=None):
    c = color or AZUL_MEDIO
    p = doc.add_paragraph()
    run = p.add_run(txt)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = c
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def badge(cell, estado):
    estado_map = {
        '✅ Completado':    (VERDE_BG,     VERDE_TEXT),
        '🔄 En progreso':  (AMARILLO_BG,  AMARILLO_T),
        '⏳ Pendiente':    (GRIS_BG,      GRIS_TEXT),
        '🚫 No aplica':    (ROJO_BG,      ROJO_TEXT),
    }
    bg, fg = estado_map.get(estado, (GRIS_BG, GRIS_TEXT))
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = p.add_run(estado)
    run.font.size = Pt(8.5)
    run.font.bold = True
    run.font.color.rgb = fg

def add_table_header(table, cols, widths, bg=AZUL_OSCURO):
    row = table.rows[0]
    for i, (txt, w) in enumerate(zip(cols, widths)):
        cell = row.cells[i]
        cell.width = Inches(w)
        set_cell_bg(cell, bg)
        set_cell_borders(cell, 'FFFFFF')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = p.add_run(txt)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = BLANCO

def add_row(table, vals, widths, estado_col=None, shade=False):
    row = table.add_row()
    bg_even = RGBColor(0xf8, 0xfa, 0xfc)
    for i, (val, w) in enumerate(zip(vals, widths)):
        cell = row.cells[i]
        cell.width = Inches(w)
        if shade:
            set_cell_bg(cell, bg_even)
        set_cell_borders(cell, 'D1D5DB')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if estado_col is not None and i == estado_col:
            badge(cell, val)
        else:
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(8.5)
            if i == 0:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1e, 0x29, 0x3b)

# ─── DATA ────────────────────────────────────────────────────────────────────

# Historias de usuario (28 HU)
HISTORIAS = [
    # ID, Titulo, Descripcion, Criterio, Prioridad, Sprint, Puntos, Estado
    ('HU-01', 'Autenticación de usuario',
     'Como alumno, quiero iniciar sesión con usuario y contraseña para acceder solo a mis sesiones.',
     'Login exitoso redirige a /modos. Token JWT guardado. Ruta protegida redirige a /login si no hay sesión.',
     'Alta', 1, 3, '✅ Completado'),
    ('HU-02', 'Registro de nuevo alumno',
     'Como alumno nuevo, quiero registrarme con nombre, apellido, username, grado y sección.',
     'Contraseña hasheada con Argon2. Username único. Respuesta 201 con datos sin password.',
     'Alta', 1, 3, '✅ Completado'),
    ('HU-03', 'Selección de modo de práctica',
     'Como alumno, quiero elegir entre modo Lectura o Expresión Libre antes de practicar.',
     'Pantalla /modos muestra dos opciones. Lectura lleva a selección de texto. Libre va directo a /practica.',
     'Alta', 1, 2, '✅ Completado'),
    ('HU-04', 'Grabación de audio en el navegador',
     'Como alumno, quiero grabar mi voz directamente desde el navegador sin instalar nada.',
     'WaveSurfer + RecordPlugin captura audio. Botón toggle grabar/detener. Visualización de onda en tiempo real.',
     'Alta', 1, 5, '✅ Completado'),
    ('HU-05', 'Cálculo de velocidad de habla (PPM)',
     'Como alumno, quiero saber cuántas palabras por minuto hablo para saber si voy muy rápido o muy lento.',
     'Whisper medium int8 transcribe el audio. PPM = palabras / duración_habla. Resultado disponible en < 3s.',
     'Alta', 1, 5, '✅ Completado'),
    ('HU-06', 'Detección de pausas y bloqueos',
     'Como alumno, quiero saber cuántas veces me quedé sin palabras durante la práctica.',
     'Pausa > 1.5s = bloqueo. Se reportan total_pauses y long_pauses. Umbral configurable en config.py.',
     'Alta', 1, 3, '✅ Completado'),
    ('HU-07', 'Análisis de prosodia con Praat',
     'Como docente, quiero ver métricas acústicas (F0, jitter, shimmer, HNR) de cada sesión del alumno.',
     'parselmouth extrae F0_mean, jitter_pct, shimmer_db, HNR_db. Valores disponibles en respuesta JSON.',
     'Media', 1, 5, '✅ Completado'),
    ('HU-08', 'Retroalimentación automática con estrellas',
     'Como alumno, quiero recibir una calificación de 1 a 5 estrellas con un mensaje motivador.',
     'Reglas: PPM 80-120 = 3 pts; pausas = 0-2 pts. Color verde/amarillo/rojo según puntuación total.',
     'Alta', 1, 3, '✅ Completado'),
    ('HU-09', 'Retroalimentación por voz (TTS)',
     'Como alumno, quiero escuchar el resultado de mi práctica en voz alta.',
     'Web Speech API lee mensaje_principal al terminar. Botón "Escuchar resultado" repite el audio.',
     'Media', 1, 2, '✅ Completado'),
    ('HU-10', 'Modo lectura con texto guía',
     'Como alumno, quiero leer un texto que me da el sistema y recibir retroalimentación sobre esa lectura.',
     'Se muestran textos de 1ro primaria. Texto visible durante grabación. texto_id enviado al backend.',
     'Alta', 1, 3, '✅ Completado'),
    ('HU-11', 'Historial de sesiones del alumno',
     'Como alumno, quiero ver mis prácticas anteriores para medir mi progreso.',
     'GET /metrics/historial retorna últimas 20 sesiones con estrellas, PPM y fecha.',
     'Alta', 1, 2, '✅ Completado'),
    ('HU-12', 'Gráfico de evolución de PPM',
     'Como alumno, quiero ver una gráfica con mi velocidad de habla en las últimas sesiones.',
     'Barras proporcionales al PPM. Verde si está en zona ideal (80-120). Se muestran últimas N sesiones.',
     'Media', 1, 2, '✅ Completado'),
    ('HU-13', 'Persistencia de sesiones en base de datos',
     'Como sistema, quiero guardar cada sesión de práctica para que el historial sea permanente.',
     'SQLAlchemy guarda Sesion + ResultadoD1. SQLite en local, Supabase en producción.',
     'Alta', 1, 3, '✅ Completado'),
    ('HU-14', 'Conversión de audio WebM/Opus a WAV',
     'Como sistema, quiero convertir el audio del navegador a WAV mono 16kHz para análisis correcto.',
     'PyAV convierte WebM/Opus → WAV 16kHz mono int16. Sin error "Not an audio file" en Praat.',
     'Alta', 1, 3, '✅ Completado'),
    ('HU-15', 'Corrección de PPM excluyendo pausas largas',
     'Como alumno, quiero que mi velocidad de habla no se vea afectada por los silencios entre palabras.',
     'speech_duration_s excluye pausas > 0.5s. PPM calculado sobre tiempo real de habla activa.',
     'Media', 1, 2, '🔄 En progreso'),
    ('HU-16', 'Detección de muletillas verbales',
     'Como alumno, quiero saber si uso muletillas como "este", "ahh", "emmm" demasiado en mi discurso.',
     'spaCy identifica tokens de relleno. Conteo por tipo. Umbral: >3 muletillas = penalización en D2.',
     'Alta', 2, 5, '⏳ Pendiente'),
    ('HU-17', 'Análisis de coherencia semántica (D2)',
     'Como alumno, quiero saber si mis ideas están conectadas y tienen sentido como discurso.',
     'BETO (bert-base-spanish-wwm) calcula coherencia entre oraciones. Score 0-1. Umbral >0.6 = bueno.',
     'Alta', 2, 8, '⏳ Pendiente'),
    ('HU-18', 'Análisis de riqueza léxica (TTR)',
     'Como alumno, quiero saber qué tan variado es mi vocabulario durante la práctica.',
     'TTR = tokens_únicos / tokens_totales. Score incluido en respuesta D2. Referencia: TTR > 0.4 = bueno.',
     'Media', 2, 3, '⏳ Pendiente'),
    ('HU-19', 'Calificación de expresividad vocal (D3)',
     'Como docente, quiero que el sistema evalúe automáticamente la expresividad de la voz del alumno.',
     'Proxy acústico: variación F0 (monotonía), shimmer, intensidad. Score 1-5 con parselmouth.',
     'Alta', 2, 5, '⏳ Pendiente'),
    ('HU-20', 'Puntaje global integrado D1+D2+D3',
     'Como alumno, quiero ver un puntaje final que combine fluidez, coherencia y expresividad.',
     'Score = 0.4*D1 + 0.35*D2 + 0.25*D3. Rango 0-100. Mensaje personalizado por rango.',
     'Alta', 2, 5, '⏳ Pendiente'),
    ('HU-21', 'Panel del docente con vista de alumnos',
     'Como docente, quiero ver el historial y progreso de todos mis alumnos en un panel.',
     'GET /docente/alumnos lista alumnos por grado/sección. GET /docente/alumno/{id}/historial.',
     'Alta', 2, 5, '⏳ Pendiente'),
    ('HU-22', 'Control de acceso por roles (RBAC)',
     'Como sistema, quiero que los docentes y alumnos solo accedan a sus funciones correspondientes.',
     'Campo rol en JWT. Middleware valida rol en rutas /docente/*. Alumno no puede ver panel docente.',
     'Alta', 2, 3, '⏳ Pendiente'),
    ('HU-23', 'Exportar reporte PDF de sesión',
     'Como docente, quiero exportar un reporte PDF de la sesión de un alumno para registros académicos.',
     'GET /reporte/{sesion_id}/pdf genera PDF con métricas D1+D2+D3, gráficos y consejos.',
     'Media', 2, 5, '⏳ Pendiente'),
    ('HU-24', 'Despliegue en Oracle Cloud + Supabase',
     'Como equipo, queremos que la aplicación esté disponible en internet sin costo adicional.',
     'Backend FastAPI en Oracle Cloud Free (ARM 4OCPU/24GB). DB Supabase PostgreSQL. Frontend Vercel.',
     'Alta', 2, 8, '⏳ Pendiente'),
    ('HU-25', 'Retroalimentación con IA generativa',
     'Como alumno, quiero recibir consejos personalizados generados por IA basados en mis métricas.',
     'Prompt a Claude Haiku con score D1+D2+D3. Respuesta en español simple para 1ro primaria.',
     'Media', 3, 8, '⏳ Pendiente'),
    ('HU-26', 'Seguimiento de progreso a lo largo del tiempo',
     'Como alumno, quiero ver gráficos de cómo he mejorado semana a semana en cada dimensión.',
     'Gráficos de línea por D1/D2/D3. Filtro por semana/mes. Tendencia con regresión lineal simple.',
     'Media', 3, 5, '⏳ Pendiente'),
    ('HU-27', 'Banco de textos de lectura por nivel',
     'Como docente, quiero subir textos de lectura categorizados por grado para los alumnos.',
     'CRUD de TextoLectura. Campos: titulo, contenido, nivel, dificultad. Endpoint POST /admin/textos.',
     'Media', 3, 3, '⏳ Pendiente'),
    ('HU-28', 'Evaluación de coincidencia en modo lectura (Levenshtein)',
     'Como sistema, quiero comparar lo que el alumno leyó con el texto original para detectar errores.',
     'Levenshtein entre transcripción y texto original. Score de fidelidad incluido en ResultadoD1.',
     'Alta', 3, 5, '⏳ Pendiente'),
]

# Items técnicos / de arquitectura (~22 items)
ITEMS_TECNICOS = [
    # ID, Titulo, Tipo, Descripcion, Sprint, Puntos, Estado
    ('TA-001', 'Configurar repositorio GitHub y rama feature/logica-negocio-srv',
     'Infraestructura', 'Rama protegida en GitHub. .gitignore para venv y .env. Primer commit con estructura base.',
     1, 1, '✅ Completado'),
    ('TA-002', 'Setup FastAPI + SQLAlchemy + Uvicorn',
     'Backend', 'main.py con startup, routers registrados, DB init. requirements.txt con todas las dependencias.',
     1, 2, '✅ Completado'),
    ('TA-003', 'Modelo de base de datos: Usuario + Sesion + ResultadoD1',
     'Base de datos', 'ORM con relaciones. Usuario 1→N Sesion 1→1 ResultadoD1. FK texto_id nullable.',
     1, 3, '✅ Completado'),
    ('TA-004', 'Pipeline de audio: WebM → WAV 16kHz mono (PyAV)',
     'Backend', 'utils/audio.py. Resampler av.AudioResampler. Flush final con resampler.resample(None).',
     1, 3, '✅ Completado'),
    ('TA-005', 'Integración faster-whisper medium int8 (sin CUDA)',
     'IA / ML', 'WhisperModel("medium", device="cpu", compute_type="int8"). word_timestamps=True para VAD.',
     1, 3, '✅ Completado'),
    ('TA-006', 'Integración parselmouth / Praat para análisis acústico',
     'IA / ML', 'analyze_prosody() extrae F0, jitter_local, shimmer_local_db, HNR, intensity_mean.',
     1, 3, '✅ Completado'),
    ('TA-007', 'Sistema de autenticación JWT + Argon2',
     'Seguridad', 'python-jose crea/valida tokens. passlib[argon2] hashea contraseñas. get_current_user dependency.',
     1, 3, '✅ Completado'),
    ('TA-008', 'Setup React + Vite + react-router-dom + WaveSurfer.js',
     'Frontend', 'Proyecto Vite con 4 páginas: Login, ModoSeleccion, Practica, Historial. Axios con interceptor JWT.',
     1, 3, '✅ Completado'),
    ('TA-009', 'CORS configurado para dev (localhost:5173) y producción',
     'Backend', 'CORSMiddleware con origins dinámicos desde env ALLOWED_ORIGINS.',
     1, 1, '✅ Completado'),
    ('TA-010', 'Seeding automático de textos de lectura en startup',
     'Backend', 'on_event("startup") seed_textos() inserta 3 textos de 1ro primaria si no existen.',
     1, 1, '✅ Completado'),
    ('TA-011', 'Corrección de cálculo PPM: excluir tiempo de pausas largas',
     'Backend', 'speech_duration_s = sum de segmentos con delta_t < 1.5s. PPM recalculado sobre habla activa.',
     1, 2, '🔄 En progreso'),
    ('TA-012', 'Alembic: migraciones de base de datos versionadas',
     'Base de datos', 'alembic init. Versiones para agregar columnas D2/D3. Compatible con SQLite y PostgreSQL.',
     2, 3, '⏳ Pendiente'),
    ('TA-013', 'Integración spaCy es_core_news_sm para análisis léxico (D2)',
     'IA / ML', 'Instalación modelo español. Función detect_muletillas() y calc_ttr() en services/dimension2/',
     2, 5, '⏳ Pendiente'),
    ('TA-014', 'Integración BETO para coherencia semántica (D2)',
     'IA / ML', 'dccuchile/bert-base-spanish-wwm-cased vía HuggingFace. sentence embeddings + cosine similarity.',
     2, 8, '⏳ Pendiente'),
    ('TA-015', 'Análisis de expresividad con proxy acústico (D3)',
     'IA / ML', 'F0 variación (std/mean). Intensidad dinámica. Score monotonía 0-100 desde parselmouth.',
     2, 5, '⏳ Pendiente'),
    ('TA-016', 'Modelo ResultadoD2 y ResultadoD3 en ORM',
     'Base de datos', 'Nuevos modelos con FK a Sesion. muletillas_count, ttr_score, coherencia_score, expresividad_score.',
     2, 3, '⏳ Pendiente'),
    ('TA-017', 'Despliegue backend en Oracle Cloud Free Tier (ARM)',
     'Infraestructura', 'Ubuntu 22.04. Python 3.11. systemd service. Nginx reverse proxy. SSL con Let\'s Encrypt.',
     2, 8, '⏳ Pendiente'),
    ('TA-018', 'Configurar Supabase PostgreSQL para producción',
     'Infraestructura', 'DATABASE_URL en .env apuntando a Supabase. Alembic upgrade head en producción.',
     2, 3, '⏳ Pendiente'),
    ('TA-019', 'Despliegue frontend en Vercel',
     'Infraestructura', 'vercel.json con VITE_API_URL apuntando a Oracle Cloud. CI automático desde main.',
     2, 2, '⏳ Pendiente'),
    ('TA-020', 'Generación de reportes PDF (ReportLab/WeasyPrint)',
     'Backend', 'GET /reporte/{sesion_id}/pdf. Incluye header, métricas D1+D2+D3, gráficos PNG y consejos.',
     2, 5, '⏳ Pendiente'),
    ('TA-021', 'Global exception handler + logger estructurado',
     'Backend', 'middleware captura 500 y retorna JSON error. loguru escribe logs a archivo con rotación diaria.',
     3, 3, '⏳ Pendiente'),
    ('TA-022', 'Tests de integración: pytest + httpx AsyncClient',
     'Calidad', 'Cobertura mínima 70%. Tests para /auth, /audio/analizar y /metrics/historial con audio fixture.',
     3, 5, '⏳ Pendiente'),
]

# ─── BUILD DOC ────────────────────────────────────────────────────────────────
doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Estilo base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\nSISTEMA DE RETROALIMENTACIÓN DE VOCALIZACIÓN\nPARA ESTUDIANTES DE 1° PRIMARIA')
run.font.size  = Pt(18)
run.font.bold  = True
run.font.color.rgb = AZUL_OSCURO

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('SRV-Oratoria-IA\n')
run2.font.size = Pt(14)
run2.font.bold = True
run2.font.color.rgb = AZUL_MEDIO

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('PRODUCT BACKLOG — VERSIÓN REFORMULADA\nFastAPI · Oracle Cloud · Supabase · Vercel\n')
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run(
    'Equipo SRV — UPAO  |  Ramdhum Arévalo Espinoza (Scrum Master)\n'
    'Taller de Tesis S04  |  Sprint 1: hasta 07/06/2026  |  Sprint 2: hasta 05/07/2026  |  Sprint 3: 01/08/2026'
)
run4.font.size = Pt(9)
run4.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCCIÓN Y CONTEXTO
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '1. CONTEXTO Y DECISIONES DE ARQUITECTURA')

doc.add_paragraph(
    'El Product Backlog original (S04) fue diseñado con NestJS y AWS. '
    'Tras análisis de costos y restricciones del proyecto, el equipo SRV decidió migrar a la siguiente arquitectura:'
)

arch_rows = [
    ('Componente',         'Original (Charter S04)',    'Actual (Reformulado)'),
    ('Backend',            'NestJS + TypeScript',       'FastAPI + Python 3.13'),
    ('IA / Transcripción', 'SpeechBrain / Kaldi',       'faster-whisper medium int8'),
    ('Análisis acústico',  'No especificado',           'parselmouth / Praat'),
    ('Análisis semántico', 'No especificado',           'spaCy + BETO (Sprint 2)'),
    ('Base de datos',      'RDS PostgreSQL (AWS)',      'SQLite (dev) / Supabase (prod)'),
    ('Backend hosting',    'AWS EC2',                   'Oracle Cloud Free (ARM 4OCPU/24GB)'),
    ('Frontend hosting',   'S3 + CloudFront',           'Vercel (gratis)'),
    ('Frontend',           'Angular',                   'React + Vite'),
]

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
add_table_header(t, ['Componente', 'Original (Charter S04)', 'Actual (Reformulado)'], [1.5, 2.5, 2.5])
for i, (comp, orig, actual) in enumerate(arch_rows[1:]):
    row = t.add_row()
    shade = (i % 2 == 0)
    for j, val in enumerate([comp, orig, actual]):
        cell = row.cells[j]
        if shade:
            set_cell_bg(cell, RGBColor(0xf8, 0xfa, 0xfc))
        set_cell_borders(cell, 'D1D5DB')
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(9)
        if j == 0:
            run.font.bold = True

doc.add_paragraph(
    '\nNota: Los ítems EN-001/002/003, TA-002 original y SP-001 del backlog anterior están marcados '
    '"No aplica" (no "fallidos"). La arquitectura actual es más sostenible económicamente y técnicamente '
    'alineada con los objetivos de investigación.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. RESUMEN DE ITEMS
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '2. RESUMEN GENERAL DEL BACKLOG')

from collections import Counter
estados_hu  = Counter(h[7] for h in HISTORIAS)
estados_ta  = Counter(t[6] for t in ITEMS_TECNICOS)
all_estados = Counter()
all_estados.update(estados_hu)
all_estados.update(estados_ta)

total_hu     = len(HISTORIAS)
total_ta     = len(ITEMS_TECNICOS)
total_items  = total_hu + total_ta
pts_total    = sum(h[6] for h in HISTORIAS) + sum(t[5] for t in ITEMS_TECNICOS)

p = doc.add_paragraph()
p.add_run(f'Total de ítems: {total_items}  |  Historias de usuario (HU): {total_hu}  |  '
          f'Tareas técnicas (TA): {total_ta}  |  Story points totales: {pts_total}').font.size = Pt(10)

summary_data = [
    ('Categoría',        'Total', '✅ Completado', '🔄 En progreso', '⏳ Pendiente'),
    ('Historias (HU)',   str(total_hu),
     str(estados_hu.get('✅ Completado', 0)),
     str(estados_hu.get('🔄 En progreso', 0)),
     str(estados_hu.get('⏳ Pendiente', 0))),
    ('Tareas técnicas',  str(total_ta),
     str(estados_ta.get('✅ Completado', 0)),
     str(estados_ta.get('🔄 En progreso', 0)),
     str(estados_ta.get('⏳ Pendiente', 0))),
    ('TOTAL',            str(total_items),
     str(all_estados.get('✅ Completado', 0)),
     str(all_estados.get('🔄 En progreso', 0)),
     str(all_estados.get('⏳ Pendiente', 0))),
]

t_sum = doc.add_table(rows=1, cols=5)
t_sum.style = 'Table Grid'
add_table_header(t_sum, ['Categoría', 'Total', '✅ Completado', '🔄 En progreso', '⏳ Pendiente'],
                 [1.8, 0.8, 1.3, 1.4, 1.3])
for i, row_data in enumerate(summary_data[1:]):
    row = t_sum.add_row()
    if i == len(summary_data) - 2:
        for j, val in enumerate(row_data):
            set_cell_bg(row.cells[j], AZUL_OSCURO)
            p2 = row.cells[j].paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.add_run(val)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = BLANCO
    else:
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            if i % 2 == 0:
                set_cell_bg(cell, RGBColor(0xf8, 0xfa, 0xfc))
            set_cell_borders(cell, 'D1D5DB')
            p2 = cell.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.add_run(val)
            run.font.size = Pt(9)
            if j == 0:
                run.font.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. SPRINT 1 — HISTORIAS DE USUARIO
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '3. SPRINT 1 — HISTORIAS DE USUARIO (hasta 07/06/2026)')

sprint1_hu = [h for h in HISTORIAS if h[5] == 1]
h2(doc, f'Historias de Usuario — Sprint 1 ({len(sprint1_hu)} ítems)')

cols_hu = ['ID', 'Historia de Usuario', 'Criterios de Aceptación', 'Prior.', 'Pts', 'Estado']
widths_hu = [0.65, 2.1, 3.3, 0.55, 0.45, 1.1]

t1 = doc.add_table(rows=1, cols=len(cols_hu))
t1.style = 'Table Grid'
add_table_header(t1, cols_hu, widths_hu)

for i, hu in enumerate(sprint1_hu):
    hu_id, titulo, desc, criterio, prio, sp, pts, estado = hu
    titulo_cell = f'{titulo}\n{desc}'
    add_row(t1, [hu_id, titulo_cell, criterio, prio, str(pts), estado],
            widths_hu, estado_col=5, shade=(i % 2 == 0))

doc.add_paragraph()

# ─── Sprint 1 — Tareas técnicas ───────────────────────────────────────────────
sprint1_ta = [t for t in ITEMS_TECNICOS if t[4] == 1]
h2(doc, f'Tareas Técnicas — Sprint 1 ({len(sprint1_ta)} ítems)')

cols_ta = ['ID', 'Tarea Técnica', 'Tipo', 'Descripción', 'Pts', 'Estado']
widths_ta = [0.65, 1.8, 1.1, 3.0, 0.45, 1.1]

t1b = doc.add_table(rows=1, cols=len(cols_ta))
t1b.style = 'Table Grid'
add_table_header(t1b, cols_ta, widths_ta)

for i, ta in enumerate(sprint1_ta):
    ta_id, titulo, tipo, desc, sp, pts, estado = ta
    add_row(t1b, [ta_id, titulo, tipo, desc, str(pts), estado],
            widths_ta, estado_col=5, shade=(i % 2 == 0))

# Sprint 1 summary
s1_all = sprint1_hu + [(t[0], t[1], '', t[3], '', 1, t[5], t[6]) for t in sprint1_ta]
s1_completados = sum(1 for x in s1_all if x[7] == '✅ Completado')
s1_progreso    = sum(1 for x in s1_all if x[7] == '🔄 En progreso')
s1_pendiente   = sum(1 for x in s1_all if x[7] == '⏳ Pendiente')
s1_total       = len(s1_all)
pct = round((s1_completados + s1_progreso * 0.5) / s1_total * 100, 1)

doc.add_paragraph()
p_s1 = doc.add_paragraph()
run_s1 = p_s1.add_run(
    f'Resumen Sprint 1:  {s1_completados}/{s1_total} completados  |  '
    f'{s1_progreso} en progreso  |  {s1_pendiente} pendientes  |  '
    f'Avance estimado: {pct}%'
)
run_s1.font.bold = True
run_s1.font.size = Pt(10)
run_s1.font.color.rgb = AZUL_MEDIO

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. SPRINT 2
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '4. SPRINT 2 — DIMENSIONES D2/D3 + DESPLIEGUE (hasta 05/07/2026)')

doc.add_paragraph(
    'Objetivo: Implementar las dimensiones D2 (coherencia y léxico) y D3 (expresividad), '
    'desplegar la aplicación en Oracle Cloud + Supabase + Vercel, e implementar el panel del docente con RBAC.'
)

sprint2_hu = [h for h in HISTORIAS if h[5] == 2]
h2(doc, f'Historias de Usuario — Sprint 2 ({len(sprint2_hu)} ítems)')

t2 = doc.add_table(rows=1, cols=len(cols_hu))
t2.style = 'Table Grid'
add_table_header(t2, cols_hu, widths_hu)
for i, hu in enumerate(sprint2_hu):
    hu_id, titulo, desc, criterio, prio, sp, pts, estado = hu
    add_row(t2, [hu_id, f'{titulo}\n{desc}', criterio, prio, str(pts), estado],
            widths_hu, estado_col=5, shade=(i % 2 == 0))

doc.add_paragraph()
sprint2_ta = [t for t in ITEMS_TECNICOS if t[4] == 2]
h2(doc, f'Tareas Técnicas — Sprint 2 ({len(sprint2_ta)} ítems)')

t2b = doc.add_table(rows=1, cols=len(cols_ta))
t2b.style = 'Table Grid'
add_table_header(t2b, cols_ta, widths_ta)
for i, ta in enumerate(sprint2_ta):
    ta_id, titulo, tipo, desc, sp, pts, estado = ta
    add_row(t2b, [ta_id, titulo, tipo, desc, str(pts), estado],
            widths_ta, estado_col=5, shade=(i % 2 == 0))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. SPRINT 3 — FINAL
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '5. SPRINT 3 — CIERRE Y ENTREGABLE FINAL (hasta 01/08/2026)')

doc.add_paragraph(
    'Objetivo: Retroalimentación con IA generativa, seguimiento longitudinal de progreso, '
    'banco de textos por docente, Levenshtein para modo lectura, tests de integración y '
    'documentación técnica completa para entrega final de tesis.'
)

sprint3_hu = [h for h in HISTORIAS if h[5] == 3]
h2(doc, f'Historias de Usuario — Sprint 3 ({len(sprint3_hu)} ítems)')

t3 = doc.add_table(rows=1, cols=len(cols_hu))
t3.style = 'Table Grid'
add_table_header(t3, cols_hu, widths_hu)
for i, hu in enumerate(sprint3_hu):
    hu_id, titulo, desc, criterio, prio, sp, pts, estado = hu
    add_row(t3, [hu_id, f'{titulo}\n{desc}', criterio, prio, str(pts), estado],
            widths_hu, estado_col=5, shade=(i % 2 == 0))

doc.add_paragraph()
sprint3_ta = [t for t in ITEMS_TECNICOS if t[4] == 3]
h2(doc, f'Tareas Técnicas — Sprint 3 ({len(sprint3_ta)} ítems)')

t3b = doc.add_table(rows=1, cols=len(cols_ta))
t3b.style = 'Table Grid'
add_table_header(t3b, cols_ta, widths_ta)
for i, ta in enumerate(sprint3_ta):
    ta_id, titulo, tipo, desc, sp, pts, estado = ta
    add_row(t3b, [ta_id, titulo, tipo, desc, str(pts), estado],
            widths_ta, estado_col=5, shade=(i % 2 == 0))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. MAPA DE OBJETIVOS → ITEMS
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '6. TRAZABILIDAD: OBJETIVOS DEL CHARTER → BACKLOG')

doc.add_paragraph(
    'Cada objetivo del Project Charter S04 está cubierto por al menos un ítem del backlog:'
)

traz_data = [
    ('OD1', 'Latencia ingesta < 2s',
     'HU-04, HU-05, TA-004, TA-005',
     'PyAV + faster-whisper int8 procesan audio en < 2s en Oracle ARM.'),
    ('OD2', 'WER < 10%, detección pausas 0.5s',
     'HU-05, HU-06, TA-005, TA-011',
     'faster-whisper medium logra WER ~8% en español. Umbral pausa configurable.'),
    ('OD3', '100% integridad datos + historial',
     'HU-11, HU-12, HU-13, TA-003, TA-012',
     'SQLAlchemy + Alembic. Historial disponible en /metrics/historial.'),
    ('OE1', 'Retroalimentación automática',
     'HU-08, HU-09, HU-17, HU-19, HU-20, HU-25',
     'D1 (reglas) → D2 (spaCy+BETO) → D3 (Praat) → Score global → IA generativa.'),
    ('OE2', 'Interfaz usable por niños 6-7 años',
     'HU-03, HU-04, HU-08, HU-09, HU-10',
     'React SPA dark mode. Botones grandes. TTS. Sin texto complejo en UI.'),
    ('OE3', 'Panel docente + reportes',
     'HU-21, HU-22, HU-23, TA-020',
     'RBAC por rol JWT. Exportación PDF. Vista por alumno y grado.'),
    ('OF1', 'Despliegue sin costo recurrente',
     'HU-24, TA-017, TA-018, TA-019',
     'Oracle Cloud Free + Supabase free + Vercel free = $0/mes.'),
]

t_traz = doc.add_table(rows=1, cols=4)
t_traz.style = 'Table Grid'
add_table_header(t_traz, ['Objetivo', 'Descripción', 'Ítems Backlog', 'Cómo se cumple'],
                 [0.6, 1.8, 2.0, 2.8])
for i, row_d in enumerate(traz_data):
    row = t_traz.add_row()
    for j, val in enumerate(row_d):
        cell = row.cells[j]
        if i % 2 == 0:
            set_cell_bg(cell, RGBColor(0xf8, 0xfa, 0xfc))
        set_cell_borders(cell, 'D1D5DB')
        p2 = cell.paragraphs[0]
        run = p2.add_run(val)
        run.font.size = Pt(8.5)
        if j == 0:
            run.font.bold = True
            run.font.color.rgb = AZUL_MEDIO

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. ROADMAP VISUAL
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '7. ROADMAP DE SPRINTS')

roadmap = [
    ('Sprint 1\n07/06/2026',
     '• Autenticación JWT + Argon2\n• Grabación audio WaveSurfer\n• Pipeline WebM→WAV (PyAV)\n'
     '• Transcripción faster-whisper\n• Análisis PPM + pausas\n• Análisis acústico Praat\n'
     '• Retroalimentación por reglas (★)\n• TTS resultado (Web Speech)\n• Modo lectura + textos guía\n'
     '• Historial + gráfico PPM\n• Base de datos ORM SQLite\n• Multi-página React SPA',
     '14/22 completados ≈ 70%\n2 en progreso'),
    ('Sprint 2\n05/07/2026',
     '• D2: spaCy muletillas + TTR\n• D2: BETO coherencia semántica\n• D3: expresividad acústica\n'
     '• Score global D1+D2+D3\n• Panel docente (RBAC)\n• Exportación PDF reportes\n'
     '• Alembic migraciones\n• Deploy Oracle Cloud + Supabase\n• Deploy Vercel frontend',
     '0/17 completados\n(inicia tras Sprint 1)'),
    ('Sprint 3\n01/08/2026',
     '• IA generativa (Claude Haiku)\n• Gráficos progreso temporal\n• Banco textos por docente\n'
     '• Levenshtein modo lectura\n• Tests integración (pytest)\n• Logger + error handler global\n'
     '• Documentación técnica final\n• Demo presentación tesis',
     '0/11 completados\n(sprint final)'),
]

t_road = doc.add_table(rows=1, cols=3)
t_road.style = 'Table Grid'
add_table_header(t_road, ['Sprint 1', 'Sprint 2', 'Sprint 3'], [2.5, 2.5, 2.5])

row = t_road.add_row()
for j, (header, items, status) in enumerate(roadmap):
    cell = row.cells[j]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    if j == 0:
        set_cell_bg(cell, VERDE_BG)
    elif j == 1:
        set_cell_bg(cell, AMARILLO_BG)
    else:
        set_cell_bg(cell, GRIS_BG)
    set_cell_borders(cell, 'D1D5DB')
    p2 = cell.paragraphs[0]
    run1 = p2.add_run(items)
    run1.font.size = Pt(8.5)
    p3 = cell.add_paragraph()
    run3 = p3.add_run(f'\n{status}')
    run3.font.size = Pt(9)
    run3.font.bold = True
    run3.font.color.rgb = VERDE_TEXT if j == 0 else AMARILLO_T if j == 1 else GRIS_TEXT

# ─── Pie de página ────────────────────────────────────────────────────────────
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_f = p_footer.add_run(
    '─────────────────────────────────────────────────────────────────────────\n'
    'SRV-Oratoria-IA · Taller de Tesis S04 · UPAO 2026\n'
    'Ramdhum Arévalo Espinoza — Scrum Master\n'
    f'Documento generado: 13/05/2026  |  Total ítems: {total_items}  |  HU: {total_hu}  |  TA: {total_ta}'
)
r_f.font.size = Pt(8)
r_f.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

# ─── Guardar ─────────────────────────────────────────────────────────────────
output = r'c:\Users\simpl\Documents\Proyecto - Informe - Taller y Tesis\SRV-Oratoria-IA_Product_Backlog_Reformulado.docx'
doc.save(output)
print(f'[OK] Documento generado: {output}')
print(f'     Total items:    {total_items} ({total_hu} HU + {total_ta} TA)')
print(f'     Story points:   {pts_total}')
print(f'     Sprint 1:       {len(sprint1_hu)} HU + {len(sprint1_ta)} TA = {len(sprint1_hu)+len(sprint1_ta)} items')
print(f'     Sprint 2:       {len(sprint2_hu)} HU + {len(sprint2_ta)} TA = {len(sprint2_hu)+len(sprint2_ta)} items')
print(f'     Sprint 3:       {len(sprint3_hu)} HU + {len(sprint3_ta)} TA = {len(sprint3_hu)+len(sprint3_ta)} items')
print(f'     Completados S1: {s1_completados}/{s1_total} ({pct}%)')

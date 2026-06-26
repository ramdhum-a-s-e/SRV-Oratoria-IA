# -*- coding: utf-8 -*-
"""Rellena la plantilla del Informe Final de Proyecto Integrador (UPAO) con el
contenido real del proyecto SRV (charter + backlog + implementación).
Conserva la portada original y aplica el formato (Arial 11, interlineado 1.5)."""
import shutil
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "0"); e.set(qn("w:color"), "BFBFBF")
        borders.append(e)
    tblPr.append(borders)

SRC = "docs/INFORME DE CAPSTONE PROJECT.docx"
DST = "docs/INFORME_CAPSTONE_SRV.docx"
shutil.copy(SRC, DST)
doc = Document(DST)

TITULO = ("Sistema de Retroalimentación por Voz (SRV) basado en Inteligencia Artificial para "
          "la mejora de la fluidez de oratoria en estudiantes de primer grado de primaria de la "
          "I.E. Juan José Farfán, Lancones - Piura")


def set_text(p, txt):
    if p.runs:
        p.runs[0].text = txt
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(txt)


# 1) Reemplazar marcadores de la portada
for p in doc.paragraphs:
    t = p.text.strip()
    if "TITULO DE PROYECTO" in t.upper():
        set_text(p, TITULO)
    elif t == "ALUMNO 1":
        set_text(p, "Br. Lezcano Saavedra, Anthony")
    elif t == "ALUMNO 2":
        set_text(p, "Br. Arévalo Espinoza, Ramdhum")
    elif t == "202X":
        set_text(p, "2026")

# 2) Borrar las instrucciones de la plantilla (desde "Resumen Ejecutivo" al final)
borrar = False
for p in list(doc.paragraphs):
    if not borrar and p.text.strip().startswith("Resumen Ejecutivo"):
        borrar = True
    if borrar:
        p._element.getparent().remove(p._element)

# 3) Formato base
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)
normal.paragraph_format.line_spacing = 1.5

# 4) Helpers de contenido
def _heading(t, level, size):
    try:
        doc.add_heading(t, level=level)
    except KeyError:
        par = doc.add_paragraph()
        r = par.add_run(t); r.bold = True; r.font.size = Pt(size)
        r.font.name = "Arial"
def h1(t): _heading(t, 1, 16)
def h2(t): _heading(t, 2, 13)
def h3(t): _heading(t, 3, 12)
def p(t):
    par = doc.add_paragraph(t)
    par.paragraph_format.line_spacing = 1.5
    return par
def b(t):
    par = doc.add_paragraph("•  " + t)
    par.paragraph_format.line_spacing = 1.5
    return par
def tabla(headers, filas):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    for i, hh in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        r = c.paragraphs[0].add_run(hh); r.bold = True; r.font.size = Pt(9)
    for fila in filas:
        cells = t.add_row().cells
        for i, v in enumerate(fila):
            cells[i].text = str(v)
            for run in cells[i].paragraphs[0].runs: run.font.size = Pt(9)
    doc.add_paragraph()

# ──────────────────────────────────────────────────────────────────────────────
# 1. RESUMEN EJECUTIVO
h1("1. Resumen Ejecutivo")
p("En la I.E. Juan José Farfán de Lancones (Piura), la evaluación de la fluidez oral en primer "
  "grado de primaria depende de la percepción subjetiva del docente; en aulas de más de 25 "
  "estudiantes resulta inviable monitorear individualmente parámetros como el ritmo, las pausas "
  "o la articulación. El presente proyecto desarrolla el Sistema de Retroalimentación por Voz "
  "(SRV), una aplicación web que captura la voz del estudiante, la procesa con modelos de "
  "inteligencia artificial y entrega una retroalimentación formativa inmediata organizada en tres "
  "dimensiones: fluidez oral, vocabulario y coherencia, y expresividad vocal.")
p("La solución integra reconocimiento automático del habla (faster-whisper), análisis prosódico "
  "(Praat/parselmouth) y procesamiento de lenguaje natural en español (spaCy y BETO), sobre una "
  "arquitectura cliente-servidor (React + FastAPI + PostgreSQL/Supabase) desplegada sin costo "
  "recurrente (Railway y Vercel). El sistema calcula un puntaje global y una calificación por "
  "estrellas, genera consejos de mejora, exporta reportes en PDF y ofrece un panel para el "
  "docente con control de acceso por roles. Una decisión central del diseño es el procesamiento "
  "local de la voz y su eliminación inmediata, para proteger los datos de los menores.")
p("Como resultado se obtuvo un sistema funcional que cubre prácticamente la totalidad del "
  "backlog del producto (alrededor de 48 de 50 ítems), desplegado y validado mediante pruebas de "
  "integración automatizadas. El aporte del proyecto no radica en un nuevo algoritmo, sino en la "
  "integración del estado del arte en un contexto educativo no atendido (niños de primer grado, "
  "zona rural, idioma español) de forma ética y de bajo costo.")
p("Palabras clave: retroalimentación por voz, fluidez oral, reconocimiento del habla, Whisper, "
  "procesamiento de lenguaje natural, educación primaria, prosodia.")

# 2. ÍNDICES
h1("2. Índices")
p("Índice del documento, índice de figuras e índice de tablas: se generan automáticamente en "
  "Word mediante Referencias → Tabla de contenido / Insertar tabla de ilustraciones, a partir de "
  "los estilos de título y los rótulos de figuras y tablas de este documento.")

# 3. DESCRIPCIÓN DEL PROYECTO
h1("3. Descripción del proyecto")
h2("3.1 Datos de la institución y sector")
tabla(["Aspecto", "Detalle"], [
    ["Institución", "I.E. Juan José Farfán"],
    ["Ubicación", "Lancones, Piura - Perú (zona rural)"],
    ["Sector", "Educación pública - Educación Básica Regular (primaria)"],
    ["Población objetivo", "Estudiantes de 1.° grado de primaria (~24-25 alumnos)"],
    ["Patrocinador / Asesor", "Mg. Walter Manuel Cueva Chávez"],
])
h2("3.2 Alcance del proyecto")
p("El proyecto comprende el diseño, desarrollo y despliegue de una aplicación web de "
  "retroalimentación por voz que incluye: autenticación segura, captura y procesamiento de audio "
  "en el navegador, análisis de fluidez (velocidad, pausas y prosodia), análisis de vocabulario "
  "y coherencia, análisis de expresividad vocal, cálculo de un puntaje global, generación de "
  "consejos, reportes en PDF, historial y gráficos de progreso, y un panel para el docente con "
  "control de acceso por roles.")
h2("3.3 Objetivos")
h3("3.3.1 Objetivo general")
p("Construir un sistema de retroalimentación por voz basado en inteligencia artificial que "
  "evalúe de forma objetiva la fluidez de oratoria de los estudiantes de primer grado de la I.E. "
  "Juan José Farfán, complementando la labor del docente.")
h3("3.3.2 Objetivos específicos")
tabla(["Objetivo", "Descripción", "Indicador / meta"], [
    ["OD1", "Motor de ingesta y transporte de audio de alta fidelidad",
     "Onda en tiempo real; inicio de transcripción con baja latencia"],
    ["OD2", "Lógica algorítmica de análisis prosódico y léxico",
     "Precisión de reconocimiento (WER); detección de pausas"],
    ["OD3", "Gestión de datos y visualización formativa",
     "100% de sesiones vinculadas correctamente al perfil del alumno"],
])
h2("3.4 Justificación")
p("La retroalimentación docente en oratoria suele ser global y diferida, lo que impide que el "
  "niño corrija sus disfluencias en el momento preciso y consolida hábitos de pronunciación "
  "inexacta. Un sistema automatizado proporciona una medición objetiva y consistente, reduce el "
  "sesgo perceptivo del evaluador y ofrece a cada estudiante un laboratorio de práctica "
  "personalizada. Además, genera evidencia del progreso a lo largo del tiempo, útil tanto para el "
  "estudiante como para el docente y la investigación.")
h2("3.5 Exclusiones del proyecto")
b("Corrección gramatical o sintáctica del discurso (no se evalúa la 'corrección' del contenido).")
b("Integración con sistemas oficiales del MINEDU (p. ej. SIAGIE).")
b("Soporte para dispositivos de gama baja sin Web Audio API.")
b("Entrenamiento de modelos de IA propios desde cero (se usan modelos preentrenados).")
h2("3.6 Restricciones del proyecto")
b("Equipo de dos integrantes y plazo académico definido (dos sprints).")
b("Presupuesto limitado: uso de servicios en planes gratuitos.")
b("Cómputo sin GPU (procesamiento en CPU), lo que condiciona la latencia.")
b("Marco ético-legal estricto por tratarse de menores de edad.")
h2("3.7 Asunciones o supuestos")
b("La institución cuenta con conexión a internet y equipos con navegador moderno y micrófono.")
b("Se dispone del consentimiento informado de los padres y el asentimiento de los menores.")
b("Los estudiantes usan la aplicación de forma voluntaria y supervisada.")

# 4. METODOLOGÍA
h1("4. Metodologías de gestión del producto")
h2("4.1 Revisión de marcos metodológicos")
tabla(["Marco", "Características", "Pertinencia para el proyecto"], [
    ["Scrum", "Iterativo por sprints, roles definidos, entregables incrementales",
     "Alta: permite entregar el sistema por incrementos (D1, luego D2/D3)"],
    ["Kanban", "Flujo continuo, tablero visual, límite de trabajo en progreso",
     "Alta: tablero en Jira para seguimiento de HU y TA"],
    ["XP (Extreme Programming)", "Prácticas técnicas intensivas (TDD, integración continua)",
     "Media: se adoptan pruebas e integración continua, no el marco completo"],
])
h2("4.2 Selección del marco metodológico")
p("Se adoptó un enfoque Scrumban: la planificación se organizó en dos sprints (Scrum) y el "
  "seguimiento de las Historias de Usuario (HU) y Tareas Técnicas (TA) se gestionó con un tablero "
  "Kanban en Jira. Se incorporaron prácticas de XP como las pruebas de integración automatizadas "
  "y la integración continua mediante GitHub.")
h2("4.3 Plan de desarrollo")
tabla(["Sprint", "Foco", "Resultado principal"], [
    ["Sprint 1", "Infraestructura, autenticación, captura/procesamiento de audio y fluidez (D1)",
     "Arquitectura base, pipeline de audio y análisis de fluidez mecánica"],
    ["Sprint 2", "Vocabulario (D2), expresividad (D3), puntaje global, panel docente, despliegue y cierre",
     "Sistema completo desplegado con las tres dimensiones y reportes"],
])
p("El producto se organizó en 8 épicas (EP-01 a EP-08) que agrupan 28 Historias de Usuario y 22 "
  "Tareas Técnicas (50 ítems, ~190 puntos de historia).")

# 5. FACTIBILIDAD Y VIABILIDAD
h1("5. Estudio de factibilidad y viabilidad")
h2("5.1 Estudio de factibilidad")
tabla(["Dimensión", "Análisis"], [
    ["Técnica", "Viable: tecnologías maduras y de código abierto (FastAPI, faster-whisper, Praat, spaCy/BETO, React). Ejecución en CPU."],
    ["Legal", "Viable con consentimiento informado, asentimiento de menores y supervisión del comité de ética; procesamiento local y audio efímero."],
    ["Organizacional", "Viable: equipo de dos integrantes con roles definidos (Project Manager y Scrum Master)."],
    ["Operativa", "Viable: aplicación web accesible desde el aula; uso voluntario y supervisado por el docente."],
    ["Recursos humanos", "Dos estudiantes de Ingeniería de Sistemas e IA; asesoría del docente patrocinador."],
    ["Ambiental", "Impacto bajo: solución de software; sin hardware especializado de alto consumo."],
    ["Riesgos", "Ruido del aula, conectividad rural, WER en voz infantil y límites de memoria en despliegue (ver mitigaciones)."],
])
h3("5.1.1 Riesgos y mitigaciones")
tabla(["Riesgo", "Mitigación"], [
    ["Ruido ambiental del aula rural", "Filtros de audio (paso-alto + puerta de ruido) y micrófonos cardioides."],
    ["Conectividad limitada", "Aplicación web ligera; modelo de procesamiento por lotes."],
    ["WER elevado en voz infantil", "Modelo Whisper 'medium' en español y medición/validación del WER."],
    ["Límite de memoria en despliegue", "PyTorch CPU-only y dimensionamiento del modelo según el plan."],
    ["Datos sensibles de menores", "Procesamiento local, audio efímero y consentimiento informado."],
])
h2("5.2 Estudio económico")
p("Al reformular la arquitectura hacia servicios gestionados en planes gratuitos (Railway, "
  "Supabase y Vercel) y reconocimiento del habla local (en lugar de un servicio en la nube de "
  "pago), el costo operativo recurrente se reduce prácticamente a cero.")
h3("5.2.1 CAPEX (inversión inicial)")
tabla(["Concepto", "Importe (S/)"], [
    ["Micrófonos de diadema (cancelación de ruido) x25", "800.00"],
    ["Suministros (cables, adaptadores USB)", "100.00"],
    ["Papelería y material de oficina", "40.00"],
    ["Total CAPEX", "940.00"],
])
h3("5.2.2 OPEX (operación mensual)")
tabla(["Concepto", "Importe (S/)"], [
    ["Hosting backend (Railway - plan gratuito/Hobby)", "~0.00"],
    ["Base de datos (Supabase - plan gratuito)", "0.00"],
    ["Frontend (Vercel - plan gratuito)", "0.00"],
    ["API de IA generativa (opcional, por uso)", "~0.00 (crédito mínimo cubre la fase de pruebas)"],
])
h3("5.2.3 Análisis costo-beneficio")
p("Por tratarse de un proyecto de investigación académica sin flujo de ingresos, los indicadores "
  "VAN y TIR no son aplicables en sentido financiero estricto. El beneficio principal es "
  "cualitativo: reducción de la subjetividad en la evaluación, ahorro de horas de evaluación "
  "manual del docente (estimado en S/ 500 por su valor de tiempo) y generación de evidencia de "
  "progreso. La relación costo-beneficio es favorable dado el bajo costo de operación.")

# 6. DESARROLLO
h1("6. Desarrollo del proyecto")
p("El desarrollo siguió el plan de dos sprints definido en el numeral 4.3, sobre la arquitectura "
  "cliente-servidor descrita en el Anexo de Documentación de Arquitectura.")
h2("6.1 Sprint 1 - Núcleo de infraestructura, audio y fluidez (D1)")
b("Autenticación segura con JWT y hashing Argon2; registro y validación de datos (DTOs).")
b("Captura de audio en el navegador con visualización de onda en tiempo real (WaveSurfer).")
b("Conversión de audio a WAV 16 kHz (PyAV) y transcripción con faster-whisper.")
b("Cálculo de velocidad de habla (PPM), detección de pausas/bloqueos y prosodia con Praat.")
b("Persistencia de sesiones y resultados en base de datos; historial del alumno.")
h2("6.2 Sprint 2 - Dimensiones D2/D3, panel docente, despliegue y cierre")
b("Detección de muletillas y riqueza léxica (TTR) con spaCy; coherencia semántica con BETO.")
b("Evaluación de expresividad vocal (variación tonal y calidad) como proxy acústico.")
b("Puntaje global ponderado (0.4·D1 + 0.35·D2 + 0.25·D3) y calificación por estrellas.")
b("Modo lectura con fidelidad por distancia de Levenshtein.")
b("Filtros anti-ruido, reportes en PDF, gráficos de progreso (Chart.js) y consejos.")
b("Control de acceso por roles (RBAC) y panel del docente.")
b("Consejos con IA generativa (Claude Haiku) de forma opcional, enviando solo métricas numéricas.")
b("Despliegue en Railway, Supabase y Vercel; pruebas de integración automatizadas.")

# 7. RESULTADOS
h1("7. Resultados")
p("Se obtuvo un sistema web funcional y desplegado que implementa las tres dimensiones de "
  "análisis y la totalidad del flujo de retroalimentación. Respecto al backlog del producto (50 "
  "ítems), se completaron alrededor de 48; las dos tareas de infraestructura previstas para AWS y "
  "migraciones se sustituyeron por alternativas equivalentes (Railway/Supabase y una "
  "micro-migración propia).")
tabla(["Indicador", "Resultado"], [
    ["Ítems del backlog cumplidos", "~48 de 50"],
    ["Dimensiones de análisis", "3 (Fluidez, Vocabulario/Coherencia, Expresividad)"],
    ["Pruebas de integración", "28 pruebas automatizadas (cobertura ~71%)"],
    ["Reportes", "Exportación PDF por sesión, con branding institucional"],
    ["Seguridad", "JWT + Argon2 + RBAC (alumno/docente)"],
    ["Despliegue", "Backend (Railway), BD (Supabase), Frontend (Vercel)"],
])
p("Respecto a los objetivos: la visualización de la onda es en tiempo real (OD1) y la integridad "
  "de datos sesión-perfil se cumple al 100% (OD3). La meta de WER < 10% (OD2) queda sujeta a una "
  "medición con datos de uso real, dado que la recolección se realiza a través del uso de la "
  "aplicación por los estudiantes.")

# 8. CONCLUSIONES
h1("8. Conclusiones")
b("Es factible construir un sistema de retroalimentación por voz, ético y desplegable a bajo "
  "costo, que evalúe objetivamente la fluidez oral de niños de primer grado en español.")
b("La integración de tecnologías del estado del arte (Whisper, BETO y Praat) permitió cubrir las "
  "tres dimensiones de evaluación sin necesidad de entrenar modelos propios.")
b("El procesamiento local y la eliminación del audio constituyen una ventaja ética frente a las "
  "soluciones que dependen de servicios de reconocimiento en la nube.")
b("El uso de un esquema de puntuación basado en reglas resulta adecuado ante la ausencia de un "
  "conjunto de datos etiquetado de voz infantil.")
b("El sistema cumple de forma sustancial el alcance comprometido en el Project Charter y el "
  "Product Backlog.")

# 9. RECOMENDACIONES
h1("9. Recomendaciones")
b("Realizar la medición del WER con datos de uso real y complementarla con un baseline de "
  "reconocimiento puro leído por los investigadores.")
b("Emplear micrófonos cardioides con cancelación de ruido para reducir el WER en el aula rural.")
b("Considerar, como trabajo futuro, un modelo de aprendizaje automático para la clasificación de "
  "la fluidez una vez se disponga de datos etiquetados.")
b("Incorporar un banco de textos administrable por el docente y, opcionalmente, voces de síntesis "
  "neuronal para una retroalimentación auditiva más natural.")
b("Capacitar brevemente a los docentes en el uso del panel y la interpretación de las métricas.")

# 10. REFERENCIAS
h1("10. Referencias Bibliográficas")
for r in [
 "Cevallos Correa, F. L., & Gómez Ríos, M. D. (2021). Reconocimiento automático de voz aplicado a la mejora en el proceso de aprendizaje de lectura en nivel escolar (Trabajo de grado). Universidad Politécnica Salesiana, Guayaquil, Ecuador.",
 "García Pazos, E. A., Escalante Vega, J. E., Alonso Ramírez, O., & Castañeda Sánchez, F. (2025). Software de reconocimiento de voz para el desarrollo de la pronunciación de inglés: Una revisión sistemática. RITI Journal, 13(29). https://doi.org/10.36825/RITI.13.29.014",
 "Haider, F., Koutsombogera, M., Conlan, O., Vogel, C., Campbell, N., & Luz, S. (2020). An Active Data Representation of Videos for Automatic Scoring of Oral Presentation Delivery Skills and Feedback Generation. Frontiers in Computer Science, 2, 1. https://doi.org/10.3389/fcomp.2020.00001",
 "Jinga, N., Anghel, A. M., Moldoveanu, F., Moldoveanu, A., Morar, A., & Petrescu, L. (2024). Overcoming Fear and Improving Public Speaking Skills through Adaptive VR Training. Electronics, 13(11), 2042. https://doi.org/10.3390/electronics13112042",
 "Sánchez, L., Morales, A., & Rodríguez, I. (2024). Speech Recognition Software as a Tool to Enhance EFL Learners' Pronunciation. Lengua y Sociedad, 23(2), 963-983. https://doi.org/10.15381/lengsoc.v23i2.26074",
 "Sonnleitner, B., Madou, T., Deceuninck, M., Theodosiou, F., & Sagaert, Y. R. (2025). Evaluation of early student performance prediction given concept drift. Computers and Education: Artificial Intelligence, 8, 100369. Elsevier.",
]:
    par = doc.add_paragraph(r); par.paragraph_format.line_spacing = 1.5

# 11. ANEXOS
h1("11. Anexos")
b("Anexo A - Código fuente: repositorio GitHub https://github.com/ramdhum-a-s-e/SRV-Oratoria-IA (y repositorio del equipo https://github.com/ALS-12321/SRV-Oratoria-IA).")
b("Anexo B - Manual de Usuario (documento MANUAL_USUARIO.docx).")
b("Anexo C - Documentación de Arquitectura (documento ARQUITECTURA.docx).")
b("Anexo D - Justificaciones de decisiones técnicas del Charter (JUSTIFICACIONES_CHARTER.docx).")
b("Anexo E - Cuadro comparativo de antecedentes (COMPARATIVA_ANTECEDENTES.docx).")
b("Anexo F - Diagramas UML y modelo de datos (ver Documentación de Arquitectura).")

doc.save(DST)
print("Generado:", DST)

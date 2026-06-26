# -*- coding: utf-8 -*-
"""Genera el Product Backlog ACTUALIZADO de SRV-Oratoria-IA con los estados reales
del desarrollo (Railway en vez de Oracle, Python 3.11, ensure_schema en vez de
Alembic, HU/TA completadas). Salida: docs/PRODUCT_BACKLOG_ACTUALIZADO.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import Counter

AZUL_OSCURO = RGBColor(0x1e, 0x29, 0x3b); AZUL_MEDIO = RGBColor(0x1e, 0x40, 0xaf)
VERDE_BG = RGBColor(0xdc, 0xfc, 0xe7); VERDE_TEXT = RGBColor(0x16, 0x65, 0x34)
AMARILLO_BG = RGBColor(0xff, 0xf9, 0xc4); AMARILLO_T = RGBColor(0x78, 0x35, 0x00)
GRIS_BG = RGBColor(0xf1, 0xf5, 0xf9); GRIS_TEXT = RGBColor(0x47, 0x55, 0x69)
ROJO_BG = RGBColor(0xfe, 0xe2, 0xe2); ROJO_TEXT = RGBColor(0x99, 0x1b, 0x1b)
BLANCO = RGBColor(0xff, 0xff, 0xff)


def rgb_hex(rgb): return rgb if isinstance(rgb, str) else f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'
def set_cell_bg(cell, rgb):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), rgb_hex(rgb))
    tcPr.append(shd)
def set_cell_borders(cell, color='D1D5DB'):
    tcPr = cell._tc.get_or_add_tcPr(); tb = OxmlElement('w:tcBorders')
    for s in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{s}'); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4'); e.set(qn('w:color'), color); tb.append(e)
    tcPr.append(tb)
def set_table_borders(t):
    tblPr = t._tbl.tblPr; b = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4'); e.set(qn('w:color'), 'D1D5DB'); b.append(e)
    tblPr.append(b)

ESTADO_MAP = {'✅ Completado': (VERDE_BG, VERDE_TEXT), '🔄 En progreso': (AMARILLO_BG, AMARILLO_T),
              '⏳ Pendiente': (GRIS_BG, GRIS_TEXT), '🚫 No aplica': (ROJO_BG, ROJO_TEXT)}

doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8); sec.left_margin = Cm(2); sec.right_margin = Cm(2)
doc.styles['Normal'].font.name = 'Calibri'; doc.styles['Normal'].font.size = Pt(10)


def h1(txt):
    p = doc.add_paragraph(); r = p.add_run(txt); r.font.size = Pt(15); r.bold = True; r.font.color.rgb = BLANCO
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), rgb_hex(AZUL_OSCURO))
    p._p.get_or_add_pPr().append(shd); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
def h2(txt):
    p = doc.add_paragraph(); r = p.add_run(txt); r.font.size = Pt(12); r.bold = True; r.font.color.rgb = AZUL_MEDIO
    p.paragraph_format.space_before = Pt(8)
def para(txt, size=10):
    p = doc.add_paragraph(); p.add_run(txt).font.size = Pt(size); return p
def header_row(table, cols, widths):
    for i, (txt, w) in enumerate(zip(cols, widths)):
        c = table.rows[0].cells[i]; c.width = Inches(w); set_cell_bg(c, AZUL_OSCURO); set_cell_borders(c, 'FFFFFF')
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER; p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = BLANCO
def fila(table, vals, widths, estado_col=None, shade=False):
    row = table.add_row()
    for i, (v, w) in enumerate(zip(vals, widths)):
        c = row.cells[i]; c.width = Inches(w); set_cell_borders(c); c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if estado_col is not None and i == estado_col:
            bg, fg = ESTADO_MAP.get(v, (GRIS_BG, GRIS_TEXT)); set_cell_bg(c, bg)
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(v); r.font.size = Pt(8.5); r.bold = True; r.font.color.rgb = fg
        else:
            if shade: set_cell_bg(c, RGBColor(0xf8, 0xfa, 0xfc))
            p = c.paragraphs[0]; r = p.add_run(v); r.font.size = Pt(8.5)
            if i == 0: r.bold = True; r.font.color.rgb = AZUL_OSCURO

# ─── DATA con estados reales (Sprint: 1 ó 2) ──────────────────────────────────
C = '✅ Completado'; P = '⏳ Pendiente'; NA = '🚫 No aplica'
HISTORIAS = [
    ('HU-01', 'Autenticación de usuario', 'Login con usuario y contraseña.', 'JWT guardado; ruta protegida redirige a /login.', 'Alta', 1, 3, C),
    ('HU-02', 'Registro de nuevo alumno', 'Registro con nombre, apellido, usuario, grado, sección.', 'Contraseña con Argon2; username único; 201 sin password.', 'Alta', 1, 3, C),
    ('HU-03', 'Selección de modo de práctica', 'Elegir Lectura o Expresión Libre.', '/modos con dos opciones; lectura lleva a selección de texto.', 'Alta', 1, 2, C),
    ('HU-04', 'Grabación de audio en el navegador', 'Grabar voz sin instalar nada.', 'WaveSurfer + RecordPlugin; onda en tiempo real.', 'Alta', 1, 5, C),
    ('HU-05', 'Cálculo de velocidad de habla (PPM)', 'Saber cuántas palabras por minuto hablo.', 'Whisper medium int8; PPM = palabras / duración_habla.', 'Alta', 1, 5, C),
    ('HU-06', 'Detección de pausas y bloqueos', 'Saber cuántas veces me quedé sin palabras.', 'Pausa larga = bloqueo; total_pauses y long_pauses.', 'Alta', 1, 3, C),
    ('HU-07', 'Análisis de prosodia con Praat', 'Ver métricas acústicas (F0, jitter, shimmer, HNR).', 'parselmouth extrae los rasgos; disponibles en JSON.', 'Media', 1, 5, C),
    ('HU-08', 'Retroalimentación automática con estrellas', 'Calificación de 1 a 5 estrellas con mensaje.', 'Reglas por PPM y pausas; color verde/amarillo/rojo.', 'Alta', 1, 3, C),
    ('HU-09', 'Retroalimentación por voz (TTS)', 'Escuchar el resultado en voz alta.', 'Web Speech API lee el mensaje; botón "Escuchar".', 'Media', 1, 2, C),
    ('HU-10', 'Modo lectura con texto guía', 'Leer un texto que da el sistema.', 'Textos de 1ro; texto_id enviado al backend.', 'Alta', 1, 3, C),
    ('HU-11', 'Historial de sesiones del alumno', 'Ver mis prácticas anteriores.', 'GET /metrics/historial con últimas sesiones.', 'Alta', 1, 2, C),
    ('HU-12', 'Gráfico de evolución de PPM', 'Ver una gráfica de mi velocidad de habla.', 'Barras por PPM; verde en zona ideal (80-120).', 'Media', 1, 2, C),
    ('HU-13', 'Persistencia de sesiones en BD', 'Guardar cada práctica.', 'SQLAlchemy guarda Sesion + ResultadoD1; Supabase en prod.', 'Alta', 1, 3, C),
    ('HU-14', 'Conversión de audio a WAV', 'Convertir el audio del navegador a WAV 16 kHz.', 'PyAV → WAV mono 16 kHz int16.', 'Alta', 1, 3, C),
    ('HU-15', 'Corrección de PPM excluyendo pausas', 'Que los silencios no afecten mi velocidad.', 'speech_duration_s excluye pausas largas.', 'Media', 1, 2, C),
    ('HU-16', 'Detección de muletillas verbales', 'Saber si uso muletillas (este, ahh, emmm).', 'spaCy identifica relleno; conteo por tipo.', 'Alta', 2, 5, C),
    ('HU-17', 'Análisis de coherencia semántica (D2)', 'Saber si mis ideas están conectadas.', 'BETO calcula coherencia entre oraciones (0-1).', 'Alta', 2, 8, C),
    ('HU-18', 'Análisis de riqueza léxica (TTR)', 'Saber qué tan variado es mi vocabulario.', 'TTR = tokens únicos / totales.', 'Media', 2, 3, C),
    ('HU-19', 'Calificación de expresividad vocal (D3)', 'Evaluar la expresividad de la voz.', 'Proxy acústico: variación F0, calidad; con parselmouth.', 'Alta', 2, 5, C),
    ('HU-20', 'Puntaje global integrado D1+D2+D3', 'Ver un puntaje final combinado.', 'Score = 0.4·D1 + 0.35·D2 + 0.25·D3 (0-100).', 'Alta', 2, 5, C),
    ('HU-21', 'Panel del docente con vista de alumnos', 'Ver historial y progreso de los alumnos.', 'GET /metrics/docente/alumnos y /alumno/{id}.', 'Alta', 2, 5, C),
    ('HU-22', 'Control de acceso por roles (RBAC)', 'Que docentes y alumnos accedan solo a lo suyo.', 'Rol en JWT; guarda require_docente en rutas docente.', 'Alta', 2, 3, C),
    ('HU-23', 'Exportar reporte PDF de sesión', 'Exportar un PDF de la sesión.', 'GET /metrics/reporte/{id} con métricas y consejos (ReportLab).', 'Media', 2, 5, C),
    ('HU-24', 'Despliegue en Railway + Supabase', 'App disponible en internet sin costo.', 'Backend en Railway (Nixpacks). DB Supabase (pooler IPv4). Frontend Vercel.', 'Alta', 2, 8, C),
    ('HU-25', 'Retroalimentación con IA generativa', 'Consejos personalizados generados por IA.', 'Claude Haiku con los scores (solo números); flag + fallback.', 'Media', 2, 8, C),
    ('HU-26', 'Seguimiento de progreso en el tiempo', 'Ver cómo mejoré semana a semana.', 'Gráficos de línea (Chart.js) por dimensión; tendencia.', 'Media', 2, 5, C),
    ('HU-27', 'Banco de textos de lectura por nivel', 'Que el docente suba textos por grado.', 'CRUD de TextoLectura; endpoint de administración.', 'Media', 2, 3, P),
    ('HU-28', 'Coincidencia en modo lectura (Levenshtein)', 'Comparar lo leído con el texto original.', 'Levenshtein entre transcripción y texto; fidelidad.', 'Alta', 2, 5, C),
]
ITEMS = [
    ('TA-001', 'Repositorio GitHub y rama de trabajo', 'Infraestructura', 'Rama protegida; .gitignore; estructura base.', 1, 1, C),
    ('TA-002', 'Setup FastAPI + SQLAlchemy + Uvicorn', 'Backend', 'main.py con startup, routers y DB init.', 1, 2, C),
    ('TA-003', 'Modelo de BD: Usuario + Sesion + ResultadoD1', 'Base de datos', 'ORM con relaciones; FK texto_id nullable.', 1, 3, C),
    ('TA-004', 'Pipeline de audio WebM → WAV 16 kHz (PyAV)', 'Backend', 'utils/audio.py con resampler y flush final.', 1, 3, C),
    ('TA-005', 'Integración faster-whisper medium int8 (CPU)', 'IA / ML', 'WhisperModel cpu int8; word_timestamps.', 1, 3, C),
    ('TA-006', 'Integración parselmouth / Praat', 'IA / ML', 'analyze_prosody(): F0, jitter, shimmer, HNR, intensidad.', 1, 3, C),
    ('TA-007', 'Autenticación JWT + Argon2', 'Seguridad', 'python-jose + passlib[argon2]; get_current_user.', 1, 3, C),
    ('TA-008', 'Setup React + Vite + WaveSurfer', 'Frontend', 'SPA con páginas; Axios con interceptor JWT.', 1, 3, C),
    ('TA-009', 'CORS para dev y producción', 'Backend', 'CORSMiddleware con origins desde env.', 1, 1, C),
    ('TA-010', 'Seeding de textos de lectura en startup', 'Backend', 'Inserta 3 textos de 1ro si no existen.', 1, 1, C),
    ('TA-011', 'Corrección de PPM (excluir pausas largas)', 'Backend', 'speech_duration_s sobre habla activa.', 1, 2, C),
    ('TA-012', 'Migraciones de esquema de BD', 'Base de datos', 'Sustituido: esquema gestionado con ensure_schema() (micro-migración, sin Alembic).', 2, 3, NA),
    ('TA-013', 'Integración spaCy para léxico (D2)', 'IA / ML', 'detect_muletillas() y calc_ttr() (es_core_news_lg).', 2, 5, C),
    ('TA-014', 'Integración BETO para coherencia (D2)', 'IA / ML', 'bert-base-spanish-wwm-cased; similitud coseno.', 2, 8, C),
    ('TA-015', 'Expresividad con proxy acústico (D3)', 'IA / ML', 'Variación F0 e intensidad; score de monotonía.', 2, 5, C),
    ('TA-016', 'Modelos ResultadoD2 y ResultadoD3', 'Base de datos', 'FK a Sesion; muletillas, ttr, coherencia, expresividad.', 2, 3, C),
    ('TA-017', 'Despliegue backend en Railway (Nixpacks)', 'Infraestructura', 'torch CPU-only para imagen ligera; build.sh predescarga modelos; HTTPS gestionado.', 2, 8, C),
    ('TA-018', 'Configurar Supabase PostgreSQL (prod)', 'Infraestructura', 'DATABASE_URL al pooler (IPv4); esquema con ensure_schema().', 2, 3, C),
    ('TA-019', 'Despliegue frontend en Vercel', 'Infraestructura', 'VITE_API_URL apuntando a Railway; CI desde main.', 2, 2, C),
    ('TA-020', 'Generación de reportes PDF (ReportLab)', 'Backend', 'GET /metrics/reporte/{id}: métricas D1+D2+D3 y consejos.', 2, 5, C),
    ('TA-021', 'Exception handler global + logging', 'Backend', 'Middleware 500 → JSON; loguru estructurado.', 2, 3, C),
    ('TA-022', 'Tests de integración (pytest)', 'Calidad', 'Cobertura ~71%; /auth, /audio/analizar, /metrics, RBAC.', 2, 5, C),
]

# ─── PORTADA ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\nPRODUCT BACKLOG — SRV ORATORIA IA'); r.font.size = Pt(18); r.bold = True; r.font.color.rgb = AZUL_OSCURO
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Sistema de Retroalimentación por Voz basado en IA  |  Versión Actualizada · 2 Sprints'); r.font.size = Pt(12); r.bold = True; r.font.color.rgb = AZUL_MEDIO
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('FastAPI · Supabase · Railway · Vercel'); r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Equipo SRV · UPAO 2026  |  Ramdhum Arévalo Espinoza (Scrum Master)  |  Actualizado: 24/06/2026'); r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)
doc.add_paragraph()

# ─── 1. CONTEXTO / STACK ──────────────────────────────────────────────────────
h1('1. Contexto y decisiones de arquitectura')
para('El backlog se reformuló respecto al Charter S04 (AWS/NestJS) hacia una arquitectura más '
     'sostenible y económica. Cambios respecto a la versión previa: el despliegue del backend pasó '
     'de Oracle Cloud a Railway, y las migraciones con Alembic se sustituyeron por una '
     'micro-migración propia (ensure_schema).')
stack = [
    ('Backend', 'NestJS + TypeScript', 'FastAPI + Python 3.11'),
    ('IA / Transcripción', 'AWS Transcribe Streaming', 'faster-whisper medium int8'),
    ('Análisis acústico', 'No especificado', 'parselmouth / Praat'),
    ('Análisis semántico', 'No especificado', 'spaCy + BETO'),
    ('Base de datos', 'RDS PostgreSQL (AWS)', 'SQLite (dev) / Supabase (prod)'),
    ('Backend hosting', 'AWS EC2', 'Railway (Nixpacks, plan Hobby)'),
    ('Frontend hosting', 'S3 + CloudFront', 'Vercel (gratuito)'),
    ('Frontend', 'Angular', 'React + Vite'),
]
t = doc.add_table(rows=1, cols=3); set_table_borders(t)
header_row(t, ['Componente', 'Original (Charter S04)', 'Actual (Reformulado)'], [1.6, 2.4, 2.6])
for i, fdat in enumerate(stack):
    fila(t, list(fdat), [1.6, 2.4, 2.6], shade=(i % 2 == 0))
doc.add_paragraph()

# ─── 2. RESUMEN ───────────────────────────────────────────────────────────────
h1('2. Resumen general del backlog')
e_hu = Counter(h[7] for h in HISTORIAS); e_ta = Counter(t2[6] for t2 in ITEMS)
tot = Counter(); tot.update(e_hu); tot.update(e_ta)
nhu, nta = len(HISTORIAS), len(ITEMS)
pts = sum(h[6] for h in HISTORIAS) + sum(t2[5] for t2 in ITEMS)
para(f'Total de ítems: {nhu+nta}  |  Historias (HU): {nhu}  |  Tareas técnicas (TA): {nta}  |  Story points: {pts}')
ts = doc.add_table(rows=1, cols=5); set_table_borders(ts)
header_row(ts, ['Categoría', 'Total', '✅ Completado', '⏳ Pendiente', '🚫 No aplica'], [1.8, 0.9, 1.4, 1.2, 1.2])
def resumen_fila(nombre, total, cnt):
    fila(ts, [nombre, str(total), str(cnt.get(C, 0)), str(cnt.get(P, 0)), str(cnt.get(NA, 0))], [1.8, 0.9, 1.4, 1.2, 1.2])
resumen_fila('Historias (HU)', nhu, e_hu)
resumen_fila('Tareas técnicas', nta, e_ta)
resumen_fila('TOTAL', nhu + nta, tot)
doc.add_paragraph()
para('Estado: prácticamente todo el alcance cumplido. Pendiente real: HU-27 (banco de textos por '
     'docente). TA-012 figura como "No aplica" porque se resolvió sin Alembic (ensure_schema).', 9)

cols_hu = ['ID', 'Historia de Usuario', 'Criterios de Aceptación', 'Pr.', 'Pts', 'Estado']
w_hu = [0.7, 2.3, 3.0, 0.5, 0.45, 1.15]
cols_ta = ['ID', 'Tarea Técnica', 'Tipo', 'Descripción', 'Pts', 'Estado']
w_ta = [0.7, 2.0, 1.1, 2.8, 0.45, 1.15]

def tabla_hu(items):
    t = doc.add_table(rows=1, cols=6); set_table_borders(t); header_row(t, cols_hu, w_hu)
    for i, hu in enumerate(items):
        fila(t, [hu[0], f'{hu[1]}\n{hu[2]}', hu[3], hu[4], str(hu[6]), hu[7]], w_hu, estado_col=5, shade=(i % 2 == 0))
def tabla_ta(items):
    t = doc.add_table(rows=1, cols=6); set_table_borders(t); header_row(t, cols_ta, w_ta)
    for i, ta in enumerate(items):
        fila(t, [ta[0], ta[1], ta[2], ta[3], str(ta[5]), ta[6]], w_ta, estado_col=5, shade=(i % 2 == 0))

# ─── 3. SPRINT 1 ──────────────────────────────────────────────────────────────
h1('3. Sprint 1 — Infraestructura, audio y fluidez D1 (hasta 07/06/2026)')
s1h = [h for h in HISTORIAS if h[5] == 1]; s1t = [t2 for t2 in ITEMS if t2[4] == 1]
h2(f'Historias de Usuario ({len(s1h)})'); tabla_hu(s1h)
doc.add_paragraph(); h2(f'Tareas Técnicas ({len(s1t)})'); tabla_ta(s1t)
para('Resumen Sprint 1: 26/26 completados (100%).', 10)

# ─── 4. SPRINT 2 ──────────────────────────────────────────────────────────────
h1('4. Sprint 2 — Dimensiones D2/D3, panel docente, despliegue y cierre (hasta 05/07/2026)')
s2h = [h for h in HISTORIAS if h[5] == 2]; s2t = [t2 for t2 in ITEMS if t2[4] == 2]
h2(f'Historias de Usuario ({len(s2h)})'); tabla_hu(s2h)
doc.add_paragraph(); h2(f'Tareas Técnicas ({len(s2t)})'); tabla_ta(s2t)
para('Resumen Sprint 2: 21/24 completados, 1 pendiente (HU-27), 1 no aplica (TA-012). Avance ~92%.', 10)

# ─── 5. TRAZABILIDAD ──────────────────────────────────────────────────────────
h1('5. Trazabilidad: objetivos del Charter → backlog')
traz = [
    ('OD1', 'Latencia de ingesta baja', 'HU-04, HU-05, TA-004, TA-005', 'PyAV + faster-whisper; onda en tiempo real. Análisis por lotes en Railway (CPU).'),
    ('OD2', 'Precisión (WER) y detección de pausas', 'HU-05, HU-06, HU-28, TA-005', 'Whisper medium; WER medido vía Levenshtein. Pendiente medición con datos reales.'),
    ('OD3', '100% integridad datos + historial', 'HU-11, HU-12, HU-13, TA-003', 'SQLAlchemy + ensure_schema(); historial en /metrics/historial.'),
    ('OE1', 'Retroalimentación automática', 'HU-08, HU-09, HU-17, HU-19, HU-20, HU-25', 'D1 (reglas) → D2 (spaCy+BETO) → D3 (Praat) → Score global → IA generativa.'),
    ('OE2', 'Interfaz usable por niños 6-7 años', 'HU-03, HU-04, HU-08, HU-09, HU-10', 'SPA React con tema infantil "Lorito"; botones grandes; TTS; señal visual de grabación.'),
    ('OE3', 'Panel docente + reportes', 'HU-21, HU-22, HU-23, TA-020', 'RBAC por rol JWT; exportación PDF; vista por alumno.'),
    ('OF1', 'Despliegue sin costo recurrente', 'HU-24, TA-017, TA-018, TA-019', 'Railway + Supabase + Vercel (planes gratuitos) ≈ $0/mes.'),
]
tt = doc.add_table(rows=1, cols=4); set_table_borders(tt)
header_row(tt, ['Objetivo', 'Descripción', 'Ítems', 'Cómo se cumple'], [0.7, 1.8, 2.0, 2.8])
for i, rd in enumerate(traz):
    fila(tt, list(rd), [0.7, 1.8, 2.0, 2.8], shade=(i % 2 == 0))

doc.add_paragraph()
pf = doc.add_paragraph(); pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pf.add_run('SRV-Oratoria-IA · Taller de Tesis S04 · UPAO 2026 · Ramdhum Arévalo Espinoza (Scrum Master) · '
               f'{nhu+nta} ítems · {pts} story points · Actualizado 24/06/2026')
r.font.size = Pt(8); r.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

doc.save('docs/PRODUCT_BACKLOG_ACTUALIZADO.docx')
print('Generado: docs/PRODUCT_BACKLOG_ACTUALIZADO.docx')
print('HU:', dict(e_hu)); print('TA:', dict(e_ta)); print('TOTAL:', dict(tot))

# -*- coding: utf-8 -*-
"""Genera el Word de tabla comparativa Antecedentes vs. Nuestra propuesta (SRV)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT

VERDE = RGBColor(0x1e, 0x9e, 0x60)

doc = Document()
base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(10)

# Orientación horizontal (tablas anchas)
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width


def set_cell(cell, text, bold=False, size=9, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def tabla(headers, filas, anchos=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        t.style = "Light Grid Accent 1"
    except Exception:
        t.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, bold=True, size=9)
    for fila in filas:
        cells = t.add_row().cells
        for i, val in enumerate(fila):
            set_cell(cells[i], val, bold=(i == 0), size=8.5)
    return t


# ── Título ──────────────────────────────────────────────────────────────────
h = doc.add_heading("Cuadro comparativo: Antecedentes vs. Nuestra propuesta (SRV)", level=0)
p = doc.add_paragraph(
    "Sistema de Retroalimentación por Voz (SRV) basado en IA para la fluidez de oratoria — "
    "I.E. Juan José Farfán, Lancones (Piura). Documento de apoyo para la exposición. "
    "Las celdas se sustentan en los seis antecedentes citados (ver Referencias); los datos de "
    "tecnología y método provienen del resumen/introducción de cada fuente.")
p.runs[0].font.size = Pt(9.5)

# ── Tabla 1: Comparativa general ──────────────────────────────────────────────
doc.add_heading("1. Comparativa general", level=1)
tabla(
    ["Estudio (referencia)", "Objetivo", "Tecnología clave (ASR/IA)",
     "Método de evaluación / scoring", "Población", "Despliegue"],
    [
        ["Cevallos & Gómez (2021) [1]",
         "Revisar el ASR para mejorar el aprendizaje de la lectura",
         "DNN + HMM (Viterbi, Baum-Welch)",
         "Revisión/análisis de algoritmos",
         "Niños escolares (español)",
         "Estudio teórico (sin app desplegada)"],
        ["Sánchez, Morales & Rodríguez (2024) [2]",
         "Evaluar el SRS como apoyo a la pronunciación (EFL)",
         "Software de reconocimiento de voz (SRS)",
         "Estudio cualitativo (notas, grabaciones, diarios)",
         "10 universitarios EFL (Colombia)",
         "Herramienta SRS de terceros en aula"],
        ["García Pazos et al. (2025) [3]",
         "Revisión sistemática de software SRS para pronunciación de inglés",
         "Comparativa de herramientas SRS",
         "Revisión sistemática de literatura",
         "EFL (meta-estudio)",
         "No aplica (revisión)"],
        ["Jinga et al. (2024) [4]",
         "Reducir el miedo y mejorar el hablar en público",
         "Realidad Virtual; análisis de voz + movimiento",
         "Métricas en tiempo real + resumen; estudio con usuarios",
         "Adultos",
         "VR inmersiva (cascos)"],
        ["Haider et al. (2020) [5]",
         "Scoring automático de habilidades de presentación oral + feedback",
         "Descriptores audiovisuales + mapas autoorganizados (SOM)",
         "Clasificación no supervisada (Active Data Representation)",
         "Adultos (presentaciones)",
         "Sistema de investigación (video)"],
        ["Sonnleitner et al. (2025) [6]",
         "Predecir el desempeño estudiantil ante 'concept drift'",
         "ML: LASSO / regresión lineal",
         "Comparación de pipelines; 2 variables 'handcrafted'",
         "Educación superior (datos LMS)",
         "Estudio de ML (no voz)"],
        ["NUESTRA PROPUESTA — SRV",
         "Retroalimentación formativa de fluidez oral (oratoria)",
         "Whisper (faster-whisper) + Praat + BETO/spaCy",
         "Scoring por reglas en 3 dimensiones + Levenshtein (WER)",
         "Niños de 1.° grado, rural (Piura, español)",
         "App web (FastAPI+React; Railway/Supabase/Vercel); audio local y efímero"],
    ])

# ── Tabla 2: Posicionamiento frente al estado del arte ────────────────────────
doc.add_heading("2. Posicionamiento frente al estado del arte", level=1)
tabla(
    ["Eje", "Estado del arte (antecedentes)", "Nuestra propuesta (SRV)", "Nivel"],
    [
        ["ASR", "DNN+HMM [1]; audiovisual+SOM [5]; SRS [2,3]",
         "Whisper (2022+) con timestamps por palabra", "A la par / por delante"],
        ["NLP semántico", "Escasamente abordado en los antecedentes",
         "BETO (BERT en español): coherencia, TTR, muletillas", "A la par (estado del arte)"],
        ["Prosodia / rasgos acústicos", "Extracción de rasgos (p. ej. OpenSMILE)",
         "Praat: F0, jitter, shimmer, HNR, intensidad", "Equivalente"],
        ["Scoring / evaluación", "Clasificación ML [5]; ML/LASSO [6]",
         "Reglas y umbrales (3 dimensiones)",
         "Por debajo — justificado: sin dataset etiquetado; [6] respalda lo simple"],
        ["Modalidad de despliegue", "VR inmersiva [4]; video [5]",
         "App web accesible", "Distinto — por viabilidad en aula rural"],
        ["Población / contexto", "Adultos, EFL, educación superior",
         "Niños de 1.° grado, rural, español", "Aporte diferencial (nicho no atendido)"],
        ["Ética de datos", "ASR en la nube / grabación retenida",
         "Procesamiento local, audio efímero (no se envía a terceros)", "Ventaja diferencial"],
    ])

# ── Comparador recomendado ────────────────────────────────────────────────────
doc.add_heading("3. ¿Con qué antecedente comparar la app?", level=1)
for txt in [
    "Comparador directo principal: Haider et al. (2020) [5] — mismo objetivo (scoring automático de presentación oral + generación de feedback).",
    "Comparador del modelo de feedback: Jinga et al. (2024) [4] — feedback en tiempo real + resumen final (tú en web, ellos en VR).",
    "Comparador de población/contexto: Cevallos & Gómez (2021) [1] — ASR + niños escolares + lectura + español.",
    "Sustento del enfoque simple: Sonnleitner et al. (2025) [6] — modelos simples superan a ML complejo con pocos datos.",
    "Encuadre del estado del arte: Sánchez et al. (2024) [2] y García Pazos et al. (2025) [3] — revisiones de herramientas SRS.",
]:
    doc.add_paragraph(txt, style="List Bullet")

# ── Puntos clave para la exposición ───────────────────────────────────────────
doc.add_heading("4. Puntos clave para la exposición", level=1)
for txt in [
    "El núcleo tecnológico (Whisper + BETO + Praat) está al nivel del estado del arte; en ASR es más moderno que [1] (2021) y [5] (2020).",
    "El aporte no es superar un algoritmo, sino trasladar el estado del arte a una población no atendida (niños de 1.° grado, rural, español) de forma ética y desplegable.",
    "El scoring por reglas se justifica por la ausencia de un dataset etiquetado de voz infantil (no hay permiso de grabación) y se respalda en [6].",
    "Ventaja ética: el audio de menores se procesa localmente y se elimina, sin enviarse a servicios de terceros.",
]:
    doc.add_paragraph(txt, style="List Bullet")

# ── Referencias ───────────────────────────────────────────────────────────────
doc.add_heading("Referencias", level=1)
refs = [
    "[1] Cevallos Correa, F. L., & Gómez Ríos, M. D. (2021). Reconocimiento automático de voz aplicado a la mejora en el proceso de aprendizaje de lectura en nivel escolar (Trabajo de grado, Ingeniería de Sistemas). Universidad Politécnica Salesiana, Sede Guayaquil, Ecuador.",
    "[2] Sánchez, L., Morales, A., & Rodríguez, I. (2024). Speech Recognition Software as a Tool to Enhance EFL Learners' Pronunciation. Lengua y Sociedad, 23(2), 963–983. https://doi.org/10.15381/lengsoc.v23i2.26074",
    "[3] García Pazos, E. A., Escalante Vega, J. E., Alonso Ramírez, O., & Castañeda Sánchez, F. (2025). Software de reconocimiento de voz para el desarrollo de la pronunciación de inglés: Una revisión sistemática. RITI Journal, 13(29). https://doi.org/10.36825/RITI.13.29.014",
    "[4] Jinga, N., Anghel, A. M., Moldoveanu, F., Moldoveanu, A., Morar, A., & Petrescu, L. (2024). Overcoming Fear and Improving Public Speaking Skills through Adaptive VR Training. Electronics, 13(11), 2042. https://doi.org/10.3390/electronics13112042",
    "[5] Haider, F., Koutsombogera, M., Conlan, O., Vogel, C., Campbell, N., & Luz, S. (2020). An Active Data Representation of Videos for Automatic Scoring of Oral Presentation Delivery Skills and Feedback Generation. Frontiers in Computer Science, 2, 1. https://doi.org/10.3389/fcomp.2020.00001",
    "[6] Sonnleitner, B., Madou, T., Deceuninck, M., Theodosiou, F., & Sagaert, Y. R. (2025). Evaluation of early student performance prediction given concept drift. Computers and Education: Artificial Intelligence, 8, 100369. Elsevier.",
]
for r in refs:
    p = doc.add_paragraph(r)
    p.runs[0].font.size = Pt(9)

doc.save("docs/COMPARATIVA_ANTECEDENTES.docx")
print("Generado: docs/COMPARATIVA_ANTECEDENTES.docx")

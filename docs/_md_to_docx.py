"""Conversor mínimo Markdown -> .docx para los documentos del proyecto.
Soporta: # / ## / ### encabezados, **negrita**, listas '- ', listas '1. ',
citas '> ', separador '---' y tablas con pipes '| a | b |'.
Uso: python _md_to_docx.py entrada.md salida.docx
"""
import re
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_runs(paragraph, text):
    for i, frag in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        run = paragraph.add_run(frag)
        if i % 2 == 1:
            run.bold = True


def _celdas(linea):
    # "| a | b | c |" -> ["a","b","c"]
    return [c.strip() for c in linea.strip().strip("|").split("|")]


def _es_separador(linea):
    return bool(re.match(r"^\|[\s:\-\|]+\|$", linea.strip()))


def agregar_tabla(doc, filas):
    headers = _celdas(filas[0])
    cuerpo = [_celdas(f) for f in filas[2:]]  # filas[1] es el separador
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        t.style = "Light Grid Accent 1"
    except Exception:
        t.style = "Table Grid"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9)
    for fila in cuerpo:
        cells = t.add_row().cells
        for i, val in enumerate(fila):
            if i < len(cells):
                cells[i].text = ""
                add_runs(cells[i].paragraphs[0], val)
                for run in cells[i].paragraphs[0].runs:
                    run.font.size = Pt(9)


def convert(md_path, docx_path):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    with open(md_path, encoding="utf-8") as f:
        lineas = f.read().splitlines()

    i = 0
    while i < len(lineas):
        s = lineas[i].rstrip()
        if not s.strip():
            i += 1
            continue
        # Tabla: línea con pipes seguida de separador
        if s.strip().startswith("|") and i + 1 < len(lineas) and _es_separador(lineas[i + 1]):
            bloque = [s]
            j = i + 1
            while j < len(lineas) and lineas[j].strip().startswith("|"):
                bloque.append(lineas[j].rstrip())
                j += 1
            agregar_tabla(doc, bloque)
            i = j
            continue
        if s.startswith("# "):
            doc.add_heading(s[2:], level=0)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=1)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=2)
        elif s.strip() == "---":
            doc.add_paragraph()
        elif s.startswith("> "):
            add_runs(doc.add_paragraph(style="Intense Quote"), s[2:])
        elif re.match(r"^[-*] ", s):
            add_runs(doc.add_paragraph(style="List Bullet"), s[2:])
        elif re.match(r"^\d+\. ", s):
            add_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\. ", "", s))
        else:
            add_runs(doc.add_paragraph(), s)
        i += 1

    doc.save(docx_path)
    print(f"Generado: {docx_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
